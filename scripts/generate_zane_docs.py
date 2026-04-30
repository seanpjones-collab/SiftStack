"""Generate the four Acquisitions Word documents in one shot.

Outputs:
  - Acquisitions_SOP.docx                  (full daily SOP with power hours)
  - Acquisitions_Companion_Training.docx   (objections, VMs, SMS, email)
  - Acquisitions_Cheat_Sheet.docx          (single-page printable reference)
  - Acquisitions_Onboarding_Checklist.docx (Day 1 / Week 1 / Month 1 / Month 3)

Branding constants are at the top — change once, regenerate all three.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

REPO = Path(__file__).resolve().parent.parent
OUTPUT_DIR = REPO / "output" / "sift_setup"

# ── Branding ──────────────────────────────────────────────────────────
COMPANY = "Alworth Homes"
# Two numbers — both spell ALWORTH on phone keypad (A=2, L=5, W=9, O=6, R=7, T=8, H=4):
PHONE = "(844) 259-6784"           # Toll-free, primary published / branding (callbacks)
LOCAL_PHONE = "(567) 259-6784"     # Local OH, used for outbound voice + SMS (toll-free
                                   # SMS verification still pending with smrtPhone)
WEBSITE = "alworthhomes.com"

# ── Style helpers ─────────────────────────────────────────────────────


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 0:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def add_para(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text, *, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_callout(doc, text, color="warn"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Light Grid Accent 1" if color == "info" else "Light Grid Accent 2"
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True


def add_quote(doc, text, *, label=None):
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_table(doc, headers, rows, *, font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(font_size + 1)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)


def add_objection_pair(doc, objection, response):
    p = doc.add_paragraph()
    run = p.add_run("They say: ")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run = p.add_run(f'"{objection}"')
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("You say: ")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run = p.add_run(f'"{response}"')
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()


def add_pagebreak(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# DOC 1: SOP
# ─────────────────────────────────────────────────────────────────


def build_sop():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Cover
    add_heading(doc, "Acquisitions SOP", level=0)
    add_para(doc, f"{COMPANY}  |  {PHONE}  (toll-free)  |  {LOCAL_PHONE}  (local OH)  |  {WEBSITE}",
             italic=True, size=11)
    add_para(doc, "Lead Manager + Closer + Prospector (Specialist / Blueprint C)",
             italic=True, size=11)
    add_para(doc, "Markets: Cuyahoga, Summit, Stark — Probate + Foreclosure",
             italic=True, size=11)
    add_para(doc, "Last updated 2026-04-28", italic=True, size=10)
    doc.add_paragraph()

    # Phone-number explainer
    add_heading(doc, "Which Number to Use When", level=2)
    add_para(doc,
             f"You have two phone numbers, both spell ALWORTH on a phone keypad:")
    add_bullet(doc, f"{PHONE} — toll-free vanity. THIS is the one you SAY in voicemails, "
                    f"emails, and during calls when leaving a callback number. It's the "
                    f"easy-to-remember brand line.")
    add_bullet(doc, f"{LOCAL_PHONE} — local Ohio number. THIS is the line your outbound "
                    f"calls and SMS go through. It's text-compliant; the toll-free isn't yet "
                    f"(smrtPhone verification pending).")
    add_para(doc, "Why it matters: when you text a seller, the message comes from "
                  f"{LOCAL_PHONE}. When you leave a voicemail, you say to call you back "
                  f"at {PHONE} (the easier vanity). Sellers will see two numbers — that's "
                  f"fine, both ring you, both spell ALWORTH.", italic=True)
    doc.add_paragraph()

    # The 5 non-negotiables
    add_heading(doc, "The 5 Non-Negotiable Rules", level=1)
    add_callout(doc,
                "Break any of these and the entire system breaks. If something below "
                "conflicts with the rest of this SOP, the rule wins.",
                color="warn")
    doc.add_paragraph()
    rules = [
        ("Every lead must have a next step.",
         "If a lead exists in Sift without a scheduled task, it's already dying. "
         "Set the next-step task BEFORE moving to the next call. No exceptions."),
        ("Call back within 1 minute on any new inbound inquiry.",
         "Harvard data: 400% close rate uplift vs 5+ minutes. Drop everything for new inbound."),
        ("Update Property Status after every contact.",
         "Status Accuracy is Foundational (STABM). Wrong status fires wrong sequences. "
         "Status update is part of the call, not a 'later' task."),
        ("Calls between 8 AM and 9 PM only.",
         "TCPA compliance window. Even on Saturday on-call, no dialing outside this window."),
        ("Empathy first, transactional second.",
         "Probate leads are grieving. Foreclosure leads are being sued. The 'Are you "
         "interested in selling?' opener gets you hung up on. Use the niche-specific "
         "scripts later in this SOP."),
    ]
    for title, body in rules:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(title + " ")
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run = p.add_run(body)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    # POWER HOURS
    add_heading(doc, "Power Hours — Highest-Value Time Blocks", level=1)
    add_para(doc,
             "Power hours are when the most people are home and answering the phone. "
             "Results-Driven data confirms these windows are the highest-connect times. "
             "DO NOT use power hours for follow-up admin, SMS responses, or status "
             "cleanup — those happen between blocks. Power hour = pure outbound dialing.")
    doc.add_paragraph()
    power_hours = [
        ("Morning power hour", "9:00 – 10:00 AM", "Tue, Wed, Thu, Fri",
         "Catch sellers before work / first thing"),
        ("Evening power hour", "5:00 – 6:00 PM", "Mon, Tue, Wed, Thu, *Fri (see note)*",
         "Catch sellers as they get home"),
    ]
    add_table(doc, ["Block", "Time", "Days", "Why"], power_hours)
    doc.add_paragraph()
    add_callout(doc,
                "FRIDAY EVENING POWER HOUR GAP: The current Friday schedule ends at 5 PM, "
                "missing the 5–6 PM power hour. Two options: extend Friday to 6 PM (gains "
                "the highest-value hour of the week) OR drop Friday evening power hour. "
                "Sean to decide.",
                color="warn")

    # STABM morning routine
    add_pagebreak(doc)
    add_heading(doc, "STABM — 5-Minute Morning Routine", level=1)
    add_para(doc,
             "Before any calls each day, run STABM. Five layers, one quick check each. "
             "Catches every lead that would otherwise fall through the cracks.")
    stabm = [
        ("S — Status", "Every record's Property Status reflects current reality. No record "
                       "left in 'New Lead' weeks after first contact."),
        ("T — Tasks", "Every overdue task gets handled or rescheduled today. Zero overdue "
                      "is the daily target."),
        ("A — Assign", "Any unassigned records get assigned to a rep. Unassigned records "
                       "fall through the cracks because no one owns them."),
        ("B — Board", "SiftLine card position matches the property status. If the card is "
                      "on 'Hot Lead' phase but status is still 'New Lead', the wrong "
                      "sequence fires. Sync them."),
        ("M — Messages", "Review message-board comments on records you'll touch today. "
                         "Context for the call before you dial."),
    ]
    add_table(doc, ["Layer", "What to check"], stabm)
    doc.add_paragraph()
    add_callout(doc,
                "Ty's data: operators who add the 5-minute STABM routine go from losing "
                "3-4 deals a quarter to zero lost deals. Five minutes. Don't skip it.",
                color="info")

    # Schedule + power hour overlay
    add_pagebreak(doc)
    add_heading(doc, "Daily Routine (with Power Hour Overlay)", level=1)
    add_para(doc, "Power hours marked with ⚡. During those windows: no admin, no follow-up "
                  "calls — pure outbound to fresh hot leads. Each day's start-of-day block "
                  "begins with the STABM check (see prior section).", italic=True)

    # Monday
    add_heading(doc, "Monday — 2:00 PM – 8:00 PM (6 hours, 1 power hour)", level=2)
    add_para(doc, "Daily target: 50–80 dials, ~15–25 record-attempts.")
    monday = [
        ("2:00 – 2:30 PM", "Start of day",
         "Run STABM (5 min — see prior section). Then run preset '02. Ready to Call' "
         "(fresh FTM) and '00. Hot Active Calls (Legacy)'. Note total queue count — "
         "flag your manager if >40."),
        ("2:30 – 5:00 PM", "Block 1 — outbound + processing (2.5 hrs)",
         "Priority: Hot temp + 1–2-day cadence → Fresh FTM → Warm → Cold. "
         "smrtPhone click-to-dial. Status + next-step task after every call."),
        ("⚡ 5:00 – 6:00 PM", "EVENING POWER HOUR (1 hr) — pure outbound",
         "Hottest hour of Monday. Stop admin, stop replies. Dial fresh hot + warm leads "
         "back-to-back. Quick status update mid-call only."),
        ("6:00 – 6:30 PM", "Admin break",
         "Eat. Reply to texts/emails that came in during power hour. "
         "Run '03. FTM Follow Up 1' for second-attempt records."),
        ("6:30 – 7:30 PM", "Block 2 — offers + appointments (1 hr)",
         "Make offers on qualified leads. Schedule appointments. Higher-touch work."),
        ("7:30 – 8:00 PM", "End of day",
         "Final pass: every record touched today has correct status + next-step task. "
         "Send daily KPI to Sean: dials, RPC, offers, contracts."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], monday)

    # Tuesday
    add_heading(doc, "Tuesday — 9:00 AM – 6:00 PM (9 hours, 2 power hours)", level=2)
    add_para(doc, "Daily target: 80–100 dials, ~25–35 record-attempts. Highest-volume day.")
    tuesday = [
        ("⚡ 9:00 – 10:00 AM", "MORNING POWER HOUR (1 hr) — pure outbound",
         "Open Sift, glance at hot/fresh queue, then dial. No admin in this hour. "
         "Goal: 25–35 dials in this hour alone."),
        ("10:00 AM – 12:00 PM", "Block 1 — continued outbound (2 hrs)",
         "Continue priority order. Update status + tasks after each call. Target another 30–40 dials."),
        ("12:00 – 1:00 PM", "Lunch + SMS triage",
         "Eat. Reply to texts/emails. Note inbounds that need callback after lunch."),
        ("1:00 – 2:00 PM", "Admin + follow-ups (1 hr)",
         "Run '03. FTM Follow Up 1'. Process inbound replies. Update overdue tasks."),
        ("2:00 – 4:30 PM", "Block 2 — offers + appointments (2.5 hrs)",
         "Make offers on qualified leads. Push on appointments. "
         "Slower pace, deeper conversations."),
        ("4:30 – 5:00 PM", "EOD prep",
         "Final status pass on records touched. Set tomorrow's tasks. KPI snapshot."),
        ("⚡ 5:00 – 6:00 PM", "EVENING POWER HOUR (1 hr) — pure outbound",
         "Hottest hour of Tuesday. Hit Hot leads for cadence calls. "
         "Stop admin. Pure dialing."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], tuesday)

    # Wednesday
    add_heading(doc, "Wednesday — 9:00 AM – 1:00 PM, then 3:00 PM – 8:00 PM "
                     "(9 hrs split, 2 power hours)", level=2)
    add_para(doc, "Daily target: 70–90 dials, ~22–30 record-attempts.")
    wednesday = [
        ("⚡ 9:00 – 10:00 AM", "MORNING POWER HOUR (1 hr) — pure outbound",
         "Dive straight into dialing. Quick task review while phone is dialing the next number."),
        ("10:00 AM – 1:00 PM", "Block 1 — continued outbound (3 hrs)",
         "Hot + Fresh priority. Status + next-step after each call. Target 35–45 dials."),
        ("1:00 – 3:00 PM", "BREAK (off the clock)",
         "Off the clock. Don't dial. Don't update Sift. Recharge."),
        ("3:00 – 3:15 PM", "Re-orient",
         "Quick scan: any inbounds during break needing 1-min callback? Otherwise, "
         "review afternoon tasks."),
        ("3:15 – 5:00 PM", "Block 2 — follow-ups (1.75 hrs)",
         "Run '03–05. FTM Follow Up' presets. Second/third-attempt focus. SMS cleanup."),
        ("⚡ 5:00 – 6:00 PM", "EVENING POWER HOUR (1 hr) — pure outbound",
         "Hottest hour of Wednesday. Hit hot list. Pure dialing."),
        ("6:00 – 6:30 PM", "Quick admin + dinner",
         "Reply to messages, dinner break."),
        ("6:30 – 7:30 PM", "Block 3 — offers + appointments (1 hr)",
         "Make offers. Push on contract appointments. Deeper touch."),
        ("7:30 – 8:00 PM", "End of day",
         "Final pass + KPI to Sean."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], wednesday)

    # Thursday
    add_heading(doc, "Thursday — 9:00 AM – 1:00 PM, then 3:00 PM – 8:00 PM "
                     "(9 hrs split, 2 power hours)", level=2)
    add_para(doc, "Same structure as Wednesday. Daily target: 70–90 dials, ~22–30 attempts. "
                  "Follow Wednesday's table above — schedule is identical.", italic=True)

    # Friday
    add_heading(doc, "Friday — 9:00 AM – 5:00 PM (8 hours, 1 power hour as scheduled)", level=2)
    add_callout(doc,
                "Schedule reminder: Friday currently ends at 5 PM, which means you "
                "miss the 5–6 PM evening power hour. Sean is deciding whether to "
                "extend Friday or skip the evening power hour. Until decided, work "
                "the morning power hour only on Friday and aim for hard close-out by 5 PM.",
                color="warn")
    doc.add_paragraph()
    add_para(doc, "Daily target: 60–80 dials, ~20–30 attempts. Wrap up the week clean.")
    friday = [
        ("⚡ 9:00 – 10:00 AM", "MORNING POWER HOUR (1 hr) — pure outbound",
         "Dive straight in. Hot + Fresh priority."),
        ("10:00 AM – 12:00 PM", "Block 1 — continued outbound (2 hrs)",
         "Continue priority order. Target 25–35 dials in this block."),
        ("12:00 – 12:30 PM", "Lunch + SMS triage",
         "Quick lunch. Reply to texts. Note Friday-afternoon callbacks."),
        ("12:30 – 4:00 PM", "Block 2 — long block, follow-ups + offers (3.5 hrs)",
         "Friday afternoon is when sellers commit or push to next week — get specific "
         "with appointment confirms for next week. Push hard on offers."),
        ("4:00 – 5:00 PM", "EOD + week-end reporting",
         "Final status pass. Send Sean: weekly KPIs (total dials, RPC, offers, "
         "contracts, conversion ratios). Flag anything that needs Sean's attention "
         "over the weekend."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], friday)

    # Saturday
    add_heading(doc, "Saturday — 9:00 AM – 1:00 PM + on-call (4 hr base + on-call)", level=2)
    add_para(doc, "Daily target: 30–50 dials during base hours, plus all inbound.")
    saturday = [
        ("9:00 – 9:30 AM", "Start of day",
         "Run STABM (5 min). Saturday queue: hot leads (1–2-day cadence), inbound that "
         "came in overnight, appointment confirms for the coming week."),
        ("9:30 AM – 12:30 PM", "Inbound + hot lead focus (3 hrs)",
         "Saturdays are good for catching W-2 sellers home. Hit your hot list. "
         "Don't blast cold — Saturday is for high-quality outreach only."),
        ("12:30 – 1:00 PM", "End of base shift + on-call handoff",
         "Final status pass. Confirm Sift is up to date. Switch to on-call mode."),
        ("After 1:00 PM", "ON-CALL",
         "Respond to inbound calls + SMS as they come in. New inbound = call back "
         "within 1 minute, even on weekend. If no inbound, no action needed."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], saturday)

    add_heading(doc, "Sunday — Off", level=2)
    add_para(doc, "Recharge. Don't dial. Don't update Sift unless something urgent. "
                  "Inbound + emergency only.", italic=True)

    # Touch sequence
    add_pagebreak(doc)
    add_heading(doc, "What Counts as One 'Attempt'", level=1)
    add_para(doc,
             "Per Ty's niche curriculum: an 'attempt' is NOT just one phone dial. "
             "An attempt is a complete touch sequence — every channel hit on the same day:")
    sequence = [
        ("1. Email", "Sent via Sift drip OR manually if no drip configured. Day 1 only."),
        ("2. SMS", "After call attempts, same day. Niche-specific message (probate vs foreclosure)."),
        ("3. Call all phones on the record", "Every number on the record gets dialed once."),
        ("4. Voicemail (if no answer)", "Niche-specific VM script (Day 1 / Day 2 / Day 3 versions)."),
        ("5. Mailer trigger (Day 1 only)", "Handwritten letter dispatched ($1.75)."),
    ]
    add_table(doc, ["Step", "What it is"], sequence)
    doc.add_paragraph()
    add_para(doc, "All 5 steps = ONE attempt. Sift's Call Attempts counter increments by 1 "
                  "after the full sequence is done — NOT after each individual dial.")
    add_para(doc, "Three attempts (Day 1, Day 2, Day 3) = full 72-hour blitz. After 3 attempts "
                  "without contact, the lead transitions out of the call queue and into the "
                  "monthly mail rotation (preset 06 → 07).", bold=True)

    # 4 Pillars
    add_pagebreak(doc)
    add_heading(doc, "The 4 Pillars Qualification", level=1)
    add_para(doc, "Apply on EVERY contact. Lead is qualified when all 4 are answered. "
                  "3 of 4 = follow up. 2 or fewer = cold lead.")
    pillars = [
        ("Reason", "Why are they thinking about it?",
         "Specific event: probate, divorce, job loss, tired landlord, code violation"),
        ("Timeline", "When do they need this resolved?",
         "“ASAP” / “within 30 days” / before auction date / before next mortgage payment"),
        ("Condition", "What's the property like?",
         "Honest about needing work / vacant / inherited as-is"),
        ("Price", "What number are they thinking?",
         "Realistic range OR “open to offers” OR an explicit need (not asking price)"),
    ]
    add_table(doc, ["Pillar", "What to ask", "Hot signal"], pillars)
    doc.add_paragraph()
    add_para(doc, "When all 4 are answered → make the offer that day. Don't slow-play hot leads.",
             bold=True)

    # Probate script
    add_pagebreak(doc)
    add_heading(doc, "Probate Script", level=1)
    add_callout(doc,
                "Context: They likely just lost a parent or close family member. They are "
                "NOT in 'selling mode' yet. Lead with empathy or you get hung up on.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "The Opener", level=2)
    add_quote(doc,
              f'"Hi, is this [PR/Executor first name]? My name is [your name] with {COMPANY}. '
              f'First — I noticed you\'re handling [decedent\'s first name]\'s estate, '
              f'and I just wanted to say I\'m sorry for your loss. I work with families '
              f'navigating inherited properties in [Cuyahoga / Summit / Stark] County, '
              f'and I reach out because most people aren\'t sure what their options are. '
              f'Do you have a couple of minutes? I won\'t take long."')

    add_heading(doc, "If they engage", level=2)
    add_para(doc, "Don't ask 'are you selling.' Ask what they're dealing with:")
    add_quote(doc,
              '"What\'s the situation with the property — are you keeping it, renting '
              'it out, or thinking about selling once probate clears?"')
    doc.add_paragraph()
    add_para(doc, "Three buckets of response:")
    add_bullet(doc, '"We don\'t know yet" → Education mode. Explain probate timeline + options. No pressure. Set 30-day follow-up.')
    add_bullet(doc, '"We want to keep it" → "Got it. If anything changes, mind if I reach out in a few months?" Tag, schedule 90-day follow-up.')
    add_bullet(doc, '"We\'re going to sell" → 4 Pillars on the spot. Most likely candidates for an offer.')

    add_heading(doc, "Discovery questions if going to sell (in order)", level=2)
    add_numbered(doc, '"Has the probate court appointed you as the executor / personal rep yet?" — determines authority to sell')
    add_numbered(doc, '"Has anyone else looked at the property — agents, family, other investors?" — competition check')
    add_numbered(doc, '"What kind of shape is it in? Honestly — anything you\'ve been putting off?" — condition')
    add_numbered(doc, '"If everything was easy and you got a fair price, what number would feel right to you?" — price (NOT "asking price")')

    add_heading(doc, "Make the offer (when 4 Pillars confirmed)", level=2)
    add_quote(doc,
              '"Based on what you told me — [property condition + their situation], '
              'I think I can help. I can make you a cash offer with no agent fees, '
              'no repairs, and we can close as fast as the court will let us. My number '
              'would be in the [$X to $Y] range. If that\'s in the ballpark, we can talk '
              'specifics. If it\'s too low, no hard feelings — I just want to be '
              'straight with you."')

    add_heading(doc, "NEVER say to a probate lead", level=2)
    add_bullet(doc, '"Are you interested in selling?" — transactional in a grief moment, hang-up trigger')
    add_bullet(doc, '"Sorry to bother you at a difficult time" — don\'t apologize for existing')
    add_bullet(doc, '"I see your dad just passed" — using "dad" feels invasive. Use the decedent\'s first name.')
    add_bullet(doc, '"I can give you a great deal" — skeptical, salesy')

    # Foreclosure script
    add_pagebreak(doc)
    add_heading(doc, "Foreclosure Script", level=1)
    add_callout(doc,
                "Context: They're being sued by their lender. Mailbox stuffed with notices, "
                "calls from creditors, possibly facing losing their home. They've heard 100 "
                "'we buy houses' pitches. Lead with help, not pressure.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "The Opener", level=2)
    add_quote(doc,
              f'"Hi, is this [first name]? [your name] here with {COMPANY}. I\'m reaching out '
              f'because I work with homeowners in [Cuyahoga / Summit / Stark] who are '
              f'dealing with the foreclosure process, and I noticed your situation. '
              f'Look — I\'m not calling to pitch you anything right now. I just want to '
              f'make sure you know what your options are, because most people don\'t get '
              f'told. Got a couple minutes?"')

    add_heading(doc, "If they engage", level=2)
    add_para(doc, "Don't push selling. Ask what they want:")
    add_quote(doc,
              '"First question — are you trying to keep the house and get caught up, '
              'or are you ready to be done with it?"')
    doc.add_paragraph()
    add_bullet(doc, '"Trying to keep it" → "Have you talked to a HUD-approved housing counselor? They\'re free." Send link. Tag `keeping-house`. Follow up in 30 days.')
    add_bullet(doc, '"Ready to be done" → 4 Pillars. Likely qualified.')
    add_bullet(doc, '"I don\'t know" → Discovery mode. Don\'t push.')

    add_heading(doc, "Discovery questions for 'ready to be done'", level=2)
    add_numbered(doc, '"When\'s the auction date — do you have one set?" — timing pressure')
    add_numbered(doc, '"How much do you owe on it, roughly?" — equity check')
    add_numbered(doc, '"What\'s the place look like inside — any work it needs?" — condition')
    add_numbered(doc, '"What would you NEED to walk away whole and start fresh?" — price as their NEED, not asking price')

    add_heading(doc, "Make the offer (when qualified)", level=2)
    add_quote(doc,
              '"Based on what you told me, here\'s what I can do — I can pay off your '
              'loan and put [$X] in your pocket. No agent fees, no repairs, you take '
              'what you want and leave the rest. I can close before the auction date. '
              'Does that get you what you need?"')

    add_heading(doc, "NEVER say to a foreclosure lead", level=2)
    add_bullet(doc, '"I see you\'re in foreclosure" — don\'t lead with their shame')
    add_bullet(doc, '"Don\'t worry, I can save you" — savior complex')
    add_bullet(doc, '"I can give you a great deal" — you\'re getting their house cheap; don\'t gloat')
    add_bullet(doc, '"How much are you asking?" — they have no asking price; ask what they NEED')

    # Deal analysis workflow
    add_pagebreak(doc)
    add_heading(doc, "Deal Analysis — Between the Qualifying Call and the Offer Call", level=1)
    add_para(doc, "After a qualifying call where all 4 Pillars are confirmed, you don't "
                  "immediately make the offer. You take a first pass at the numbers, owner "
                  "reviews and approves (or revises), then you go back to the seller with the "
                  "approved offer.")
    add_para(doc, "The point of you doing the analysis first: it builds your judgment over time, "
                  "and it makes the owner's review faster than starting from scratch every deal.",
             italic=True)

    add_heading(doc, "Step 1 — Pull comps in SiftMap / Sift Properties", level=2)
    add_bullet(doc, "Filter to the property's ZIP and similar attributes (beds/baths/sqft within ~10%, age within 15 years)")
    add_bullet(doc, "Pull recent SOLD comps (last 6 months ideally, 12 if thin) AND active listings for trend")
    add_bullet(doc, "Note finish levels: which comps are fully renovated, partially updated, or original condition? What did each tier command per sqft?")

    add_heading(doc, "Step 2 — Determine As-Is value AND/OR ARV", level=2)
    add_bullet(doc, "As-Is: what this property would fetch in its current condition (per the seller's account)")
    add_bullet(doc, "ARV: what it would fetch fully renovated to market-tier finishes")
    add_bullet(doc, "Use the project's real-estate-comping skill (Two-Bucket ARV methodology) for the framework")

    add_heading(doc, "Step 3 — Compile repair list + rehab guesstimate", level=2)
    add_bullet(doc, "From the 4 Pillars condition question, you have the seller's account of what's needed")
    add_bullet(doc, "If you've done a walk-through, use those notes")
    add_bullet(doc, "Take a swing at a rehab budget using room-by-room ranges (kitchen, bath, flooring, HVAC, roof, exterior). Don't agonize — your guess + owner's review lands in the right zone")
    add_bullet(doc, "Use the project's rehab-estimator skill for the 4-tier framework + room ranges")

    add_heading(doc, "Step 4 — Calculate MAO (Maximum Allowable Offer)", level=2)
    add_bullet(doc, "75% rule: MAO = (ARV × 0.75) − rehab − fixed costs (closing, holding, agent if applicable)")
    add_bullet(doc, "70% rule in tighter markets or higher-risk deals (older homes, distressed neighborhoods)")
    add_bullet(doc, "Document your math so owner can audit the chain")

    add_heading(doc, "Step 5 — Bring numbers to owner for review", level=2)
    add_bullet(doc, "Send a quick summary: As-Is, ARV, repair list, rehab estimate, proposed offer ($X to $Y range)")
    add_bullet(doc, "Owner reviews, may revise the rehab estimate, may set a different MAO, may approve as-is")
    add_bullet(doc, "If owner doesn't like the numbers → re-evaluate together. Don't go back to the seller until owner approves.")

    add_heading(doc, "Step 6 — Make the offer call", level=2)
    add_bullet(doc, "Use the niche-appropriate offer script (probate or foreclosure version above)")
    add_bullet(doc, "Stay within the owner-approved range")
    add_bullet(doc, "If the seller counters outside the range → escalate to owner BEFORE agreeing to anything")

    add_callout(doc,
                "Getting better at this takes reps. Early on, expect owner to revise a lot — "
                "that's the training. After 10–20 deals through this loop, your first-pass numbers "
                "should land within 10% of the owner's revised numbers most of the time.",
                color="info")

    # Status update
    add_pagebreak(doc)
    add_heading(doc, "Status Update Playbook (after every call)", level=1)
    add_para(doc, "Use this exact mapping. No improvising.")
    statuses = [
        ("Right party, qualified, will sell", "Hot Lead", "Offer call in 1–2 days"),
        ("Right party, interested but not ready", "Warm Lead", "Follow-up in 15 days"),
        ("Right party, not motivated yet (probate)", "Cold Lead", "Re-engage in 45 days"),
        ("Right party, not motivated yet (auction)", "Cold Lead", "Re-engage in 30 days"),
        ("Right party, not motivated yet (general)", "Cold Lead", "Re-engage in 90 days"),
        ("Right party, hard no", "Not Interested", "Quarterly re-engage in 90 days"),
        ("Wrong number / dead number", "(per-phone status only — Sift auto-updates)", "Move to next number"),
        ("Right party, asked DNC", "Dead Lead + DNC tag", "Don't call again"),
        ("Already sold", "Sold", "Done"),
        ("Hostile / threatening", "Dead Lead + note", "Done"),
    ]
    add_table(doc, ["Outcome", "Property Status", "Next-step task"], statuses)
    doc.add_paragraph()
    add_callout(doc,
                "Golden rule: every lead leaves the call with a scheduled task. Even "
                "Dead Lead gets a 'verify in 6 months' if circumstances might change.",
                color="warn")

    # KPIs
    add_pagebreak(doc)
    add_heading(doc, "KPIs (what Sean watches weekly)", level=1)
    kpis = [
        ("Daily dials", "50–100 (varies by day length)"),
        ("Daily record-attempts", "20–35 (one record-attempt = full touch sequence)"),
        ("RPC per dial", "1 in 32 baseline; with Trestle scoring, target 1 in 10"),
        ("Conversations to appointment", "5:1"),
        ("Appointments to offer", "1.5:1"),
        ("Offers to contract", "4:1"),
        ("Task completion rate", "95%+"),
        ("Tasks overdue", "<5%"),
    ]
    add_table(doc, ["Metric", "Target"], kpis)
    doc.add_paragraph()
    add_para(doc,
             "Trestle phone scoring (5-tier system) drives ~4.75× connect-rate uplift. "
             "Always dial 81–100 first, then 61–80. Skip ≤40 unless every other lead is exhausted.",
             italic=True)

    # Escalation
    add_heading(doc, "When to escalate to Sean", level=1)
    add_bullet(doc, "Property under contract — handoff for transaction coordination")
    add_bullet(doc, "Seller wants negotiation beyond your authority — pull Sean into the call")
    add_bullet(doc, "Unusual situation (out-of-state heirs, divorce mid-sale, title issues) — Sean handles or routes to Deep Prospecting (data manager queue)")
    add_bullet(doc, "Anything legally weird (active eviction, code violations, tenant disputes) — escalate")
    add_bullet(doc, "Inbound that looks like press / competitor / scammer — flag and ask")

    return doc


# ─────────────────────────────────────────────────────────────────
# DOC 2: COMPANION TRAINING
# ─────────────────────────────────────────────────────────────────


def build_training():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Cover
    add_heading(doc, "Acquisitions — Companion Training", level=0)
    add_para(doc, f"{COMPANY}  |  {PHONE}  (toll-free)  |  {LOCAL_PHONE}  (local OH)  |  {WEBSITE}",
             italic=True, size=11)
    add_para(doc, "Objection Handling, Voicemails, SMS, Email Templates", italic=True, size=11)
    add_para(doc, "Companion to: Acquisitions SOP", italic=True, size=11)
    add_para(doc, "Last updated 2026-04-28", italic=True, size=10)
    doc.add_paragraph()

    add_callout(doc,
                "Use this with the SOP, not instead of it. The SOP tells you WHEN to call, "
                "WHAT preset, WHICH status. This document tells you HOW to handle the call.",
                color="warn")

    # Tone
    add_heading(doc, "Tone Rules", level=1)
    add_bullet(doc, "Calm, low-pressure. Distressed sellers are sensitive to anyone who sounds urgent or hungry.")
    add_bullet(doc, "Speak slower than naturally. Drops anxiety on their end.")
    add_bullet(doc, "Pause after they finish a sentence. Most people fill silence with more information — that's where qualifying details come from.")
    add_bullet(doc, "Mirror their language. They say 'the house' → don't say 'the property.' They say 'my mom' → don't say 'the deceased.'")
    add_bullet(doc, "Never argue. If they push back, agree, ask one more question, then back off.")
    add_bullet(doc, "Smile when you talk. They can hear it.")

    # Objections — Probate
    add_pagebreak(doc)
    add_heading(doc, "Objection Handling — Probate", level=1)
    add_objection_pair(doc,
        "We're not selling.",
        "Totally understand — most families don't think about it right after. Mind if I "
        "send you a free probate guide so when you're ready to think about it, you have "
        "the info? No follow-up unless you ask.")
    add_objection_pair(doc,
        "We're going to keep it in the family.",
        "That's great. Most people end up keeping it. If anything changes — kids don't "
        "want it, taxes get heavy, repairs pile up — would it be okay if I checked in "
        "once a year? Just a quick text.")
    add_objection_pair(doc,
        "We have an agent already.",
        "Oh good. Just so you know your options — agents typically need 30 to 90 days "
        "and you'd need to deal with showings, repairs, and inspections. We can usually "
        "close in 2 to 3 weeks cash if that ever sounds appealing. Otherwise, I'm out "
        "of your hair.")
    add_objection_pair(doc,
        "I'm in the middle of probate. I can't sell yet.",
        "Right, the court has to appoint you first. Has that happened? Usually takes "
        "30 to 60 days. If you want, I can put your file aside and check back in a "
        "month — no obligation, just so we're ready when you are.")
    add_objection_pair(doc,
        "What's it worth?",
        "Honest answer — depends on condition. If you can spare 5 minutes, I can give "
        "you a ballpark over the phone, or I can swing by and walk through it. "
        "Whichever you prefer.")
    add_objection_pair(doc,
        "How did you find out about this?",
        "Public record — the probate filing in [county] court. I track those because "
        "most families don't realize how much help is available.")
    add_objection_pair(doc,
        "Stop calling me.",
        f"Got it, I'll take you off the list right now. If anything changes, my "
        f"number is {PHONE}. Take care.")
    add_objection_pair(doc,
        "Are you a scammer?",
        f"Fair question. I'm a local investor in Cleveland — you can verify the "
        f"company at {WEBSITE}. Take your time. I'd rather you check first than "
        f"feel pressured.")

    # Objections — Foreclosure
    add_pagebreak(doc)
    add_heading(doc, "Objection Handling — Foreclosure", level=1)
    add_objection_pair(doc,
        "I'm not selling.",
        "Got it. Just so you know — there are a few options most folks don't realize: "
        "HUD counseling (free), short sale, deed in lieu, and yes, a cash sale. "
        "Mind if I send a one-page summary so you have it for reference? No pressure.")
    add_objection_pair(doc,
        "I'm working with my bank.",
        "Good. Are they offering you a modification or asking you to sell? I ask "
        "because I see how lenders move every day, and I can usually tell if "
        "they're being straight with you. Happy to give you my honest take if it helps.")
    add_objection_pair(doc,
        "I have an attorney.",
        "Smart move. While they're working that, do you have a backup plan in case it "
        "doesn't work out? Sometimes the legal side takes longer than the auction date "
        "and people get caught.")
    add_objection_pair(doc,
        "I'll figure it out myself.",
        "Totally fair. If I can ask one thing — do you know your auction date? That's "
        "the clock that matters. Once I know that, I can tell you whether you've got "
        "time to figure it out, or whether you're cutting it close.")
    add_objection_pair(doc,
        "How did you get my information?",
        "Public record — the foreclosure filing in [county] court. I track those "
        "because most homeowners in that situation don't get told their options.")
    add_objection_pair(doc,
        "Stop calling me.",
        f"Done. I'll take you off the list. If your situation changes, my number is "
        f"{PHONE}. Hope it works out for you.")
    add_objection_pair(doc,
        "Are you one of those investors trying to lowball me?",
        "Fair to ask. I'm not going to lowball you — what I can pay depends on what "
        "you owe and what condition the place is in. If my number doesn't work for "
        "you, no hard feelings, you walk away. I just want to be a real option, not "
        "a vulture.")
    add_objection_pair(doc,
        "I just need a few more weeks.",
        "Okay. Do you have a specific plan in those few weeks — bank work-out, family "
        "loan, refinance? I ask because if it doesn't come through, the auction date "
        "doesn't move. If you want a backup option in place just in case, I can have "
        "an offer ready to go within 24 hours.")

    # Voicemails
    add_pagebreak(doc)
    add_heading(doc, "Voicemail Scripts", level=1)
    add_callout(doc,
                "Voicemails follow the 3-day cadence. Each VM is shorter than the last — "
                "by Day 3 you're respecting their silence.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "Probate Voicemails", level=2)
    add_quote(doc,
              f"Hi [first name], this is [your name] with {COMPANY}. I'm reaching out about "
              f"[decedent first name]'s estate — sorry for your loss. I help families "
              f"in [county] navigate inherited properties. No pressure on this call — "
              f"I just want to make sure you know your options when you're ready. My "
              f"number is {PHONE}. Take care.",
              label="VM 1 — Day 1 (~25 sec)")
    add_quote(doc,
              f"Hi [first name], [your name] again. Just wanted to make myself available. "
              f"If you're sorting through what to do with the property and want a "
              f"second opinion or just info on the process, I'm happy to help — no "
              f"pressure, no pitch. {PHONE}. Talk soon.",
              label="VM 2 — Day 2 (~20 sec)")
    add_quote(doc,
              f"Hi [first name], last try from me. If now's not the right time, "
              f"totally understand. If you ever want to talk about [property "
              f"address] or anything related to [decedent first name]'s estate, "
              f"I'm at {PHONE}. Take care.",
              label="VM 3 — Day 3 (~15 sec)")

    add_heading(doc, "Foreclosure Voicemails", level=2)
    add_quote(doc,
              f"Hi [first name], [your name] here with {COMPANY}. I work with homeowners "
              f"going through the foreclosure process in [county]. I'm not calling "
              f"to sell you anything — I just want to make sure you know your "
              f"options. There's more than people realize, and most are free. "
              f"{PHONE}. Take care.",
              label="VM 1 — Day 1 (~25 sec)")
    add_quote(doc,
              f"Hi [first name], [your name] again. Quick one — if you're trying to keep "
              f"the house, there's free HUD help. If you're ready to walk away, "
              f"there are options too. Either way I'm happy to walk through it "
              f"with you. {PHONE}.",
              label="VM 2 — Day 2 (~20 sec)")
    add_quote(doc,
              f"Hi [first name], one more try. If you've got an auction date "
              f"coming up, the clock matters. If you want help understanding "
              f"options before then, I'm at {PHONE}. No pressure either way.",
              label="VM 3 — Day 3 (~15 sec)")

    # SMS
    add_pagebreak(doc)
    add_heading(doc, "SMS Follow-Up Templates", level=1)
    add_callout(doc,
                "SMS goes after the call attempt + voicemail on the same day. Short, "
                "specific, no marketing language. Always sign off with your name.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "Probate SMS Sequence", level=2)
    add_quote(doc,
              f"Hi [first name], this is [your name] with {COMPANY}. I left you a voicemail "
              f"about [decedent first name]'s estate. I help families with inherited "
              f"properties — no pressure, just info if you want it. Reply anytime.",
              label="SMS 1 — Day 1, after VM")
    add_quote(doc,
              "Hi [first name], following up. If you're sorting through what to do "
              "with [property address], happy to walk through options no charge. -[your name]",
              label="SMS 2 — Day 2")
    add_quote(doc,
              f"Hi [first name], last text from me unless you reach out. My number "
              f"is {LOCAL_PHONE} if you ever want to talk about [property address]. "
              f"Take care. -[your name]",
              label="SMS 3 — Day 3")

    add_heading(doc, "Foreclosure SMS Sequence", level=2)
    add_quote(doc,
              f"Hi [first name], [your name] with {COMPANY}. I help homeowners in [county] "
              f"dealing with foreclosure. Not pitching anything — just options. "
              f"Reply if you want to know what they are.",
              label="SMS 1 — Day 1, after VM")
    add_quote(doc,
              "Hi [first name], if you're trying to keep the house, free HUD help "
              "exists. If you're ready to walk away, I can usually close before any "
              "auction date. Either way I'm here. -[your name]",
              label="SMS 2 — Day 2")
    add_quote(doc,
              f"Hi [first name], last text. Number's {LOCAL_PHONE} if you change your "
              f"mind. Wishing you the best. -[your name]",
              label="SMS 3 — Day 3")

    add_heading(doc, "Reply handling", level=2)
    add_bullet(doc, "'Stop' / 'Don't text me' / 'F off' → Reply 'Got it, taking you off. Take care.' Then Property Status = Dead Lead, add DNC tag.")
    add_bullet(doc, "'Who is this?' → 'Public record — the [foreclosure / probate] filing in [county] court. Happy to send my contact info if you want to verify.'")
    add_bullet(doc, "'Maybe / I'm thinking' → 'Happy to chat whenever. When's a good time to call?' Schedule the callback as a Sift task.")
    add_bullet(doc, "'How much would you pay?' → 'Depends on condition + what you owe. Easiest is a quick 5-min call — got 5 min today?' Schedule the call.")

    # Email
    add_pagebreak(doc)
    add_heading(doc, "Email Templates", level=1)
    add_callout(doc,
                "Email is the FIRST channel in the Pendulum (Email → SMS → Call → "
                "Mail → DP). Cheapest touch, creates a paper trail. Lead-in to the "
                "call sequence, not a substitute.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "Probate Email — Day 1", level=2)
    add_para(doc, "Subject: A quick note about [decedent first name]'s estate", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              f"Hi [first name],\n\n"
              f"I'm [your name] with {COMPANY} in Cleveland. First — I'm sorry for your "
              f"loss. I came across [decedent first name]'s estate in the probate "
              f"court records and wanted to reach out.\n\n"
              f"I work with families in [county] navigating inherited properties. "
              f"Most of the time, families aren't sure what their options are: "
              f"keep and rent, sell traditionally, or sell as-is for cash. "
              f"There's no right answer — it depends on what you and your family "
              f"want.\n\n"
              f"If you'd like to talk through it (no pitch, just information), my "
              f"direct line is {PHONE}. Or visit {WEBSITE} if you'd rather check "
              f"us out first.\n\n"
              f"Take care,\n[your name]")

    add_heading(doc, "Probate Email — Day 7 follow-up", level=2)
    add_para(doc, "Subject: RE: A quick note about [decedent first name]'s estate", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              f"Hi [first name],\n\n"
              f"Following up on my earlier note. I know there's a lot on your "
              f"plate right now, so no pressure.\n\n"
              f"If [property address] is something you're starting to think "
              f"about, I can give you a no-obligation cash offer in 48 hours. "
              f"If you'd rather list it traditionally or hold onto it, I can "
              f"point you to people who can help with that too.\n\n"
              f"Either way, my direct line is {PHONE}.\n\n"
              f"Take care,\n[your name]")

    add_heading(doc, "Foreclosure Email — Day 1", level=2)
    add_para(doc, "Subject: Quick note about [property address]", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              f"Hi [first name],\n\n"
              f"I'm [your name] with {COMPANY} in Cleveland. I came across the "
              f"foreclosure filing for [property address] in [county] court "
              f"records and wanted to reach out — not to pitch you anything, "
              f"but to make sure you know what your options are.\n\n"
              f"Most homeowners in this situation don't get told that there are "
              f"several paths forward:\n\n"
              f"1. HUD-approved housing counselors (free) help you negotiate "
              f"with the bank for a loan modification\n"
              f"2. Short sale — your bank approves a sale for less than what "
              f"you owe\n"
              f"3. Deed in lieu — give the property back without going to auction\n"
              f"4. Cash sale — sell to an investor before the auction and walk "
              f"away with cash\n\n"
              f"Each has trade-offs. Happy to walk through them with you — no "
              f"pitch, just information. Direct line: {PHONE}.\n\n"
              f"Either way, please don't let the auction date sneak up on you. "
              f"Take care,\n[your name]")

    add_heading(doc, "Foreclosure Email — Day 7 follow-up", level=2)
    add_para(doc, "Subject: RE: Quick note about [property address]", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              f"Hi [first name],\n\n"
              f"Following up. If you've already figured out a path, that's great "
              f"— I hope it works out.\n\n"
              f"If you haven't, the clock is the thing that matters most. "
              f"Auction dates don't move because you're working on it.\n\n"
              f"If you want to put a backup option in place — even just to know "
              f"the number — I can have a written cash offer ready in 24 hours. "
              f"No obligation. You compare it to whatever else you're working on.\n\n"
              f"Direct line: {PHONE}.\n\n"
              f"Take care,\n[your name]")

    return doc


# ─────────────────────────────────────────────────────────────────
# DOC 3: PRINTABLE CHEAT SHEET (1 page front + back)
# ─────────────────────────────────────────────────────────────────


def build_cheatsheet():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header
    add_heading(doc, "Acquisitions Cheat Sheet", level=0)
    add_para(doc, f"{COMPANY}  •  {PHONE}  •  {WEBSITE}", italic=True, size=11)
    doc.add_paragraph()

    # Power hours
    add_heading(doc, "POWER HOURS — pure outbound only", level=2)
    add_para(doc, "  ⚡ 9:00 – 10:00 AM   Tue, Wed, Thu, Fri", size=11)
    add_para(doc, "  ⚡ 5:00 – 6:00 PM    Mon, Tue, Wed, Thu  *(Fri ends 5pm — gap)*", size=11)
    add_para(doc, "  Daily target: 50–100 dials, 20–35 record-attempts", italic=True, size=10)

    # 5 rules
    add_heading(doc, "THE 5 RULES", level=2)
    rules = [
        "Every lead has a next-step task before you hang up.",
        "Inbound callback within 1 minute. Drop everything.",
        "Update Property Status after every call.",
        "Calls 8 AM – 9 PM only.",
        "Empathy first. Never 'are you interested in selling?'",
    ]
    for r in rules:
        add_bullet(doc, r, size=10)

    # 4 Pillars
    add_heading(doc, "4 PILLARS — qualify before offering", level=2)
    pillars_short = [
        ("Reason", "Why now?"),
        ("Timeline", "How fast?"),
        ("Condition", "What shape?"),
        ("Price", "What number / what do they NEED?"),
    ]
    add_table(doc, ["Pillar", "Ask"], pillars_short, font_size=9)
    add_para(doc, "All 4 = make offer today. 3 = follow up. 2 or fewer = cold.",
             italic=True, size=10)

    # Probate opener
    add_heading(doc, "PROBATE OPENER — memorize", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(
        f'"Hi [name]? [your name] with {COMPANY}. I noticed you\'re handling [decedent '
        f'first name]\'s estate — I\'m sorry for your loss. I work with families '
        f'navigating inherited properties in [county]. Most folks aren\'t sure '
        f'what their options are. Got a couple minutes? I won\'t take long."'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)

    # Foreclosure opener
    add_heading(doc, "FORECLOSURE OPENER — memorize", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(
        f'"Hi [name]? [your name] with {COMPANY}. I work with homeowners in [county] '
        f'dealing with the foreclosure process. Look — I\'m not calling to '
        f'pitch you anything. I just want to make sure you know what your '
        f'options are, because most people don\'t get told. Got a couple minutes?"'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)

    # ─── Page 2 ───
    add_pagebreak(doc)

    # Status mapping
    add_heading(doc, "STATUS MAPPING — after every call", level=2)
    status_short = [
        ("Qualified, will sell", "Hot Lead", "Offer in 1–2d"),
        ("Interested, not ready", "Warm Lead", "FU in 15d"),
        ("Cold (probate)", "Cold Lead", "45d"),
        ("Cold (auction-driven)", "Cold Lead", "30d"),
        ("Cold (general)", "Cold Lead", "90d"),
        ("Hard no", "Not Interested", "Quarterly 90d"),
        ("DNC requested", "Dead Lead + DNC", "Don't call"),
        ("Already sold", "Sold", "Done"),
    ]
    add_table(doc, ["Outcome", "Status", "Next-step"], status_short, font_size=9)

    # Touch sequence
    add_heading(doc, "ONE ATTEMPT = full sequence per day", level=2)
    add_para(doc, "  1. Email   2. SMS   3. Call all phones   4. VM if no answer   5. Mail (D1 only)",
             size=10)
    add_para(doc, "  3 attempts in 72 hrs = full cycle. After 3, transition to mail nurture.",
             italic=True, size=10)

    # Trestle
    add_heading(doc, "TRESTLE DIAL ORDER", level=2)
    trestle = [
        ("81–100", "Dial First — top priority"),
        ("61–80", "Dial Second"),
        ("41–60", "Dial Third"),
        ("21–40", "Dial Fourth — only if queue empty"),
        ("0–20", "Drop — don't dial"),
    ]
    add_table(doc, ["Score", "Action"], trestle, font_size=9)

    # Never say
    add_heading(doc, "NEVER SAY", level=2)
    add_bullet(doc, "Probate: 'Are you interested in selling?', 'Sorry to bother you', 'I see your dad just passed', 'I can give you a great deal'", size=10)
    add_bullet(doc, "Foreclosure: 'I see you're in foreclosure', 'Don't worry, I can save you', 'How much are you asking?', 'I can give you a great deal'", size=10)

    # Escalate
    add_heading(doc, "ESCALATE TO SEAN", level=2)
    add_bullet(doc, "Under contract — handoff to transaction coordination", size=10)
    add_bullet(doc, "Beyond your authority — pull Sean into the call", size=10)
    add_bullet(doc, "Out-of-state heirs / divorce / title issues — Sean handles or DP queue", size=10)
    add_bullet(doc, "Anything legally weird (eviction, code violations, tenants)", size=10)

    # Footer
    doc.add_paragraph()
    add_para(doc,
             f"Direct: {PHONE}  •  {WEBSITE}  •  Acquisitions Cheat Sheet — "
             f"Last updated 2026-04-28",
             italic=True, size=8)

    return doc


# ─────────────────────────────────────────────────────────────────
# DOC 4: ONBOARDING CHECKLIST
# ─────────────────────────────────────────────────────────────────


def build_process_retraining():
    """For existing reps adopting new tools/scripts/cadence — NOT a full ramp.

    Audience: someone who already calls leads competently but is migrating from
    an older workflow to the Sift-driven setup (current scripts, STABM, the
    13-preset niche sequential, Trestle scoring, etc.).
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Cover
    add_heading(doc, "Acquisitions Process Retraining Guide", level=0)
    add_para(doc, f"{COMPANY}  |  {PHONE}  |  {LOCAL_PHONE}  |  {WEBSITE}",
             italic=True, size=11)
    add_para(doc, "For existing reps moving to the current Sift-driven workflow",
             italic=True, size=11)
    add_para(doc, "Last updated 2026-04-29", italic=True, size=10)
    doc.add_paragraph()

    add_callout(doc,
                "Audience: existing acquisitions reps who already know how to call leads "
                "and are now adopting new tools, scripts, and cadence. This is NOT a "
                "from-scratch ramp — it's a migration guide for what's changing.",
                color="info")
    doc.add_paragraph()

    # What's changing
    add_heading(doc, "What's Changing", level=1)
    add_para(doc, "The four big shifts in your daily work:")

    changes = [
        ("CRM",
         "Old: Podio with custom workflows.",
         "New: REI Sift (app.reisift.io). Properties / SiftLine boards / Sequences. "
         "All historical data has been migrated."),
        ("Scripts",
         "Old: 'Are you interested in selling now or in the near future?' (Results-Driven style).",
         "New: empathy-first openers per niche. Probate leads with grief. Foreclosure leads "
         "with 'I'm not pitching you.' Generic openers get hung up on in distress niches."),
        ("Cadence",
         "Old: ad-hoc follow-ups.",
         "New: Hot 1–2d / Warm 15d / Cold 45d. Niche-specific re-engagement (Auction 30d / "
         "Probate 45d / General 90d). Touch sequence = email + SMS + call all phones + VM "
         "+ Day-1 mailer = ONE attempt."),
        ("Morning routine",
         "Old: jump in and start dialing.",
         "New: STABM 5-minute check (Status / Tasks / Assign / Board / Messages) before "
         "first call every day."),
    ]
    add_table(doc, ["Area", "Was", "Now"], changes)
    doc.add_paragraph()

    # The 4 documents
    add_heading(doc, "The 4 Reference Documents", level=1)
    add_para(doc, "All four live in the team's shared folder. Print or bookmark all four.")
    docs_table = [
        ("Acquisitions SOP", "Daily routine per day, power hours, status mapping, KPIs, "
                             "escalation rules. Read once end-to-end."),
        ("Companion Training", "Tone, 16 objection-handling pairs, voicemail scripts, SMS "
                               "templates, email templates. Reference for handling specific "
                               "situations."),
        ("Cheat Sheet", "One-page printable. Power hours, openers, status mapping, Trestle "
                        "dial order. Keep on desk."),
        ("This Retraining Guide", "What's changing — read once during your migration week, "
                                  "then archive."),
    ]
    add_table(doc, ["Document", "Use"], docs_table)

    # Migration timeline
    add_pagebreak(doc)
    add_heading(doc, "Your Migration Week", level=1)
    add_para(doc, "Aggressive but realistic. By end of week, your daily flow runs entirely "
                  "in the new system. No 4-week ramp — you already know the job.")

    add_heading(doc, "Day 1 — Read + System Access", level=2)
    add_bullet(doc, "Read the SOP cover-to-cover (~45 min). Note which sections feel new vs. familiar.")
    add_bullet(doc, "Skim the Companion Training and bookmark the objection sections that feel useful.")
    add_bullet(doc, "Print the Cheat Sheet (front + back). Tape it to your desk.")
    add_bullet(doc, "Verify Sift access works. Pull up your Properties view and confirm migrated records are visible.")
    add_bullet(doc, "Run STABM once on yesterday's records to feel the routine before tomorrow.")

    add_heading(doc, "Day 2 — Run STABM Live, Use New Openers", level=2)
    add_bullet(doc, "Start with the STABM check (5 min). It will feel slow on Day 2 — by Day 5 it's automatic.")
    add_bullet(doc, "Use the probate + foreclosure openers verbatim today. Don't paraphrase — let the new "
                    "language land on the seller before you adapt it.")
    add_bullet(doc, "Apply 4 Pillars (Reason / Timeline / Condition / Price) on every contact. "
                    "Note which pillar you tend to skip — that's your training edge.")
    add_bullet(doc, "Update Property Status + next-step task after every call.")

    add_heading(doc, "Day 3-5 — Settle In", level=2)
    add_bullet(doc, "Continue STABM mornings. Time it — should drop from 8 min Day 2 to ~4 min by Day 5.")
    add_bullet(doc, "Start improvising the openers in your own voice while keeping the empathy-first frame.")
    add_bullet(doc, "Use SMS / email follow-ups from the templates. Reply to inbound within 1 minute.")
    add_bullet(doc, "End each day with the daily KPI report to ownership (dials / RPC / offers / contracts).")

    add_heading(doc, "End of week 1 review (with owner)", level=2)
    add_bullet(doc, "Which scripts felt natural? Which felt forced?")
    add_bullet(doc, "Where did STABM catch something that would have slipped?")
    add_bullet(doc, "Any objection that came up that wasn't in the training doc → flag for owner to add")
    add_bullet(doc, "Trestle dial order working? RPC rate moving up?")

    # New behaviors checklist
    add_pagebreak(doc)
    add_heading(doc, "New Behaviors Checklist", level=1)
    add_para(doc, "Tick each as it becomes automatic. By Week 4, all should feel like default behavior, "
                  "not effort.")

    behaviors = [
        ("STABM run before any calls each day"),
        ("Empathy-first opener on every probate call"),
        ("'I'm not pitching you' opener on every foreclosure call"),
        ("4 Pillars asked on every right-party contact (RPC)"),
        ("Property Status updated immediately after every call"),
        ("Next-step task set on every record before moving on"),
        ("Trestle dial order respected (81-100 first, skip ≤40)"),
        ("Inbound callback within 1 minute"),
        ("SMS sent same-day after voicemail (one-two punch)"),
        ("Daily KPI report sent to ownership"),
        ("Recording self-review at least 1x per week"),
    ]
    for b in behaviors:
        add_bullet(doc, b)

    # In scope, with review
    add_heading(doc, "In Scope — You Do First, Owner Reviews", level=1)
    add_para(doc, "These are part of the role. You take a first pass; the owner reviews and "
                  "revises before any number goes out to a seller. The point of you doing the "
                  "analysis first is twofold: it builds your judgment over time, and it makes "
                  "the owner's review faster than starting from scratch every deal.")
    add_bullet(doc, "Comping / ARV — pull sold and listed comps yourself, identify finish levels "
                    "and what they're commanding, determine as-is value AND/OR ARV. Owner reviews.")
    add_bullet(doc, "Repair list + rehab guesstimate — collect needed repairs from the seller "
                    "(condition questions in the 4 Pillars) and walk-through if you can. "
                    "Take a swing at a rehab budget. Owner reviews and revises.")
    add_bullet(doc, "Offer construction — propose the offer number to the owner based on your "
                    "comp + rehab analysis. Owner approves before it goes to the seller. If "
                    "the owner doesn't like the number, you re-evaluate together.")

    # Out of scope (narrower)
    add_heading(doc, "Out of Scope for the Acquisitions Role", level=1)
    add_para(doc, "These genuinely sit with ownership — bring deals reaching these stages rather "
                  "than handling solo:")
    add_bullet(doc, "Transaction coordination (post-contract paperwork, title, closing)")
    add_bullet(doc, "Disposition — selling under-contract deals to end buyers")
    add_bullet(doc, "Marketing strategy / lead source decisions / data spend")

    # Red flags
    add_heading(doc, "Red Flags — Pause and Ask Owner", level=1)
    add_bullet(doc, "30+ min on a call without progress on 4 Pillars")
    add_bullet(doc, "Seller asks something you don't know — say 'let me get you a precise answer' and ask")
    add_bullet(doc, "Legally complex situation (active eviction, untenanted tenants, title clouds, divorce mid-listing)")
    add_bullet(doc, "Out-of-state heirs or PRs (probate cases spanning jurisdictions)")
    add_bullet(doc, "Seller offers something off-script (creative financing, partial sale, lease-back)")
    add_bullet(doc, "Pressure to offer above what feels right — pause, consult, respond")
    add_bullet(doc, "Threats from seller (legal, regulatory, personal) — document + escalate immediately")

    # Footer
    doc.add_paragraph()
    add_para(doc,
             "Once the new behaviors are automatic (typically 2-4 weeks), this guide can be archived. "
             "Daily reference moves to the SOP and Cheat Sheet. New objections or situations get added "
             "to the Companion Training over time.",
             italic=True, size=10)

    return doc


def build_onboarding():
    """For NEW HIRES coming into the acquisitions role.

    Trainer/supervisor = senior acquisitions rep (not the owner). Owner only
    reviews/approves the rep's comp + rehab + offer numbers, and handles
    ownership-only items (TC, disposition, marketing strategy).
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Cover
    add_heading(doc, "New Hire Onboarding Checklist — Acquisitions", level=0)
    add_para(doc, f"{COMPANY}  |  {PHONE}  |  {LOCAL_PHONE}  |  {WEBSITE}",
             italic=True, size=11)
    add_para(doc, "Day 1 → Month 3 ramp under the senior acquisitions rep",
             italic=True, size=11)
    add_para(doc, "Last updated 2026-04-29", italic=True, size=10)
    doc.add_paragraph()

    add_callout(doc,
                "Trainer / supervisor = the senior acquisitions rep. They run shadow sessions, "
                "review your recordings, give daily feedback, and sign off on milestones. "
                "Ownership reviews and approves the rep's comp + rehab + offer numbers "
                "before they go to a seller, and handles ownership-only items: transaction "
                "coordination, disposition, marketing strategy / data spend.",
                color="warn")
    doc.add_paragraph()

    # Day 1
    add_heading(doc, "Day 1 — Setup + Read", level=1)
    add_para(doc, "Goal: Get all access set up and absorb the four reference documents before any calls.",
             italic=True)
    add_heading(doc, "Access checklist (set up by ownership ahead of Day 1)", level=2)
    add_bullet(doc, "Sift / DataSift login — log in, change password, enable 2FA")
    add_bullet(doc, "smrtPhone login — log in, verify both numbers route to you (toll-free + local)")
    add_bullet(doc, "Email account ([firstname]@alworthhomes.com)")
    add_bullet(doc, "Calendar set up + connected to Sift / Google for appointment booking")
    add_bullet(doc, "Phone access (Slack / SMS / cell) for direct line to senior rep + ownership for escalations")

    add_heading(doc, "Read in this order (~3 hours)", level=2)
    add_numbered(doc, "Acquisitions SOP (full read — daily routine, scripts, status mapping, KPIs)")
    add_numbered(doc, "Companion Training (objections, VMs, SMS templates, email templates)")
    add_numbered(doc, "Cheat Sheet (memorize the openers — print and keep at desk)")
    add_numbered(doc, "Sift's own niche-sequential help article: intercom.help/reisift/en/articles/12543919")

    add_heading(doc, "Shadow the senior rep (~2 hours)", level=2)
    add_bullet(doc, "Sit with the senior rep during 5–10 live calls (probate + foreclosure mix)")
    add_bullet(doc, "Watch how they handle status updates + next-step task assignment in Sift")
    add_bullet(doc, "Observe at least one full power-hour block")
    add_bullet(doc, "Listen to 3–5 recorded calls (good ones + at least one bad one with notes from the senior rep)")

    # Week 1
    add_pagebreak(doc)
    add_heading(doc, "Week 1 — Live Calls + Daily Feedback", level=1)
    add_para(doc, "Goal: First ~50 live calls with daily feedback from the senior rep. By Friday, "
                  "you're solo on most calls and the senior rep spot-checks recordings rather than "
                  "listening live.", italic=True)

    add_heading(doc, "Daily (Mon–Fri)", level=2)
    add_bullet(doc, "Run STABM (5 min) before any calls")
    add_bullet(doc, "Run preset '02. Ready to Call' for fresh FTM — work it with the senior rep listening")
    add_bullet(doc, "Use the probate + foreclosure openers EXACTLY (no paraphrasing yet)")
    add_bullet(doc, "Apply 4 Pillars on every contact — even if it feels mechanical")
    add_bullet(doc, "Update Property Status + next-step task after every call")
    add_bullet(doc, "End-of-day: send daily KPI to senior rep (dials / RPC / offers / contracts)")
    add_bullet(doc, "End-of-day: 15-min recording review with senior rep")

    add_heading(doc, "Week 1 milestones (signed off by senior rep)", level=2)
    week1 = [
        ("By Day 2", "First RPC (right party contact) on a probate AND a foreclosure lead"),
        ("By Day 3", "Apply 4 Pillars unprompted (no senior-rep cue) on at least 3 calls"),
        ("By Day 4", "Make first cash offer to a qualified lead (with senior rep on standby)"),
        ("By Day 5", "Solo on power-hour blocks. Senior rep spot-checks recordings, not live."),
    ]
    add_table(doc, ["Milestone", "What it looks like"], week1)

    # Month 1
    add_pagebreak(doc)
    add_heading(doc, "Month 1 — Independent Operation", level=1)
    add_para(doc, "Goal: Hit baseline KPIs solo. Senior rep escalates only when needed.",
             italic=True)

    add_heading(doc, "Week 2", level=2)
    add_bullet(doc, "All call blocks solo. Daily KPI report still goes to senior rep.")
    add_bullet(doc, "Track scripted-vs-improvised ratio — by end of week, ~50/50 with all 4 Pillars hit unprompted")
    add_bullet(doc, "First contract written (senior rep walks you through paperwork; ownership signs)")

    add_heading(doc, "Week 3", level=2)
    add_bullet(doc, "KPI baseline: 50+ dials/day on full days, 30+ on split days")
    add_bullet(doc, "RPC rate climbing toward 1-in-10 with Trestle scoring")
    add_bullet(doc, "Manage your own Sift task queue without daily senior-rep intervention")
    add_bullet(doc, "Weekly KPI review with senior rep (Friday afternoon, 30 min)")

    add_heading(doc, "Week 4", level=2)
    add_bullet(doc, "Daily targets at upper end (80–100 dials on Tue/Wed/Thu)")
    add_bullet(doc, "5+ offers made in the week")
    add_bullet(doc, "First fully-solo contract closed (handoff to ownership for transaction coordination)")
    add_bullet(doc, "End-of-month review with senior rep + ownership — what's working, what to adjust")

    # Month 3
    add_pagebreak(doc)
    add_heading(doc, "Month 3 — Mastery Checkpoints", level=1)
    add_para(doc, "Goal: Run the entire acquisitions function without daily senior-rep involvement. "
                  "Senior rep reviews weekly KPIs, ownership reviews monthly.", italic=True)

    add_heading(doc, "Performance baselines (must hit by end of Month 3)", level=2)
    perf = [
        ("Daily dials", "50–100 (varies by day length per schedule)"),
        ("RPC rate", "≥1 in 10 with Trestle scoring"),
        ("Conversation → appointment", "5:1"),
        ("Appointment → offer", "1.5:1"),
        ("Offer → contract", "4:1"),
        ("Task completion", "95%+"),
        ("Tasks overdue", "<5%"),
        ("Contracts written / month", "2+ (Blueprint C target: 6–24/year)"),
    ]
    add_table(doc, ["Metric", "Target"], perf)

    add_heading(doc, "Soft-skill checkpoints", level=2)
    add_bullet(doc, "Handle unexpected objections without consulting the script")
    add_bullet(doc, "Identify motivated sellers in the first 30 seconds of a call")
    add_bullet(doc, "Negotiate offers within authority — pull senior rep / ownership only for genuine edge cases")
    add_bullet(doc, "Maintain Sift status accuracy without prompts")
    add_bullet(doc, "Inbound callback within 1 minute even during admin / breaks")

    add_heading(doc, "In scope — you do first, owner reviews", level=2)
    add_para(doc, "These are part of the role. You take a first pass; ownership reviews and "
                  "revises before any number reaches the seller:", italic=True)
    add_bullet(doc, "Comping / ARV — pull sold and listed comps yourself, identify finish levels and "
                    "what they're commanding, determine as-is and/or ARV. Owner reviews.")
    add_bullet(doc, "Repair list + rehab guesstimate — collect needed repairs from the 4 Pillars "
                    "condition questions (or walk-through if you have one) and propose a rehab "
                    "budget. Owner reviews and revises.")
    add_bullet(doc, "Offer construction — propose the offer number based on your comp + rehab "
                    "analysis. Owner approves before it goes to the seller. If owner doesn't like "
                    "the number, you re-evaluate together.")

    add_heading(doc, "Out of scope for the acquisitions role (handled by ownership)", level=2)
    add_para(doc, "These genuinely sit with ownership — bring deals reaching these stages rather "
                  "than handling solo:", italic=True)
    add_bullet(doc, "Transaction coordination (post-contract paperwork, title, closing)")
    add_bullet(doc, "Disposition — selling under-contract deals to end buyers")
    add_bullet(doc, "Marketing strategy / lead source decisions / data spend")

    # Continuous learning
    add_heading(doc, "Continuous learning (every week, indefinitely)", level=1)
    add_bullet(doc, "Listen to 1 of your own recorded calls per week — note what worked, what didn't")
    add_bullet(doc, "Read 1 case study from datasift.ai/case-studies — see how other operators handle similar situations")
    add_bullet(doc, "Weekly KPI review with senior rep (30 min Friday)")
    add_bullet(doc, "Monthly script-tuning session — what objections came up that aren't in the training doc?")
    add_bullet(doc, "Quarterly: read the latest Ty content for any cadence/script updates")

    # Red flags
    add_pagebreak(doc)
    add_heading(doc, "Red Flags — Pause and Escalate", level=1)
    add_para(doc, "Escalate to the senior rep first. They escalate to ownership for ownership-level decisions.",
             italic=True)
    add_bullet(doc, "30+ min on a call without progress on 4 Pillars")
    add_bullet(doc, "Seller asks a question you don't know — say 'let me get you a precise answer' and check with senior rep")
    add_bullet(doc, "Legally complex situation (active eviction, untenanted tenants, title clouds, divorce mid-listing)")
    add_bullet(doc, "Out-of-state heirs or PRs (probate cases spanning jurisdictions)")
    add_bullet(doc, "Seller offers something off-script (creative financing, partial sale, lease-back)")
    add_bullet(doc, "Pressure to offer above what feels right — pause, consult, respond")
    add_bullet(doc, "Seller threatens (legal, regulatory, personal) — document + escalate immediately")

    # Footer
    doc.add_paragraph()
    add_para(doc,
             "If something isn't covered in any of the four documents (SOP, Companion Training, "
             "Cheat Sheet, this Onboarding Checklist) — that's a signal we need to write it down. "
             "Tell the senior rep and we'll add it.",
             italic=True, size=10)

    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for name, builder in [
        ("Acquisitions_SOP.docx", build_sop),
        ("Acquisitions_Companion_Training.docx", build_training),
        ("Acquisitions_Cheat_Sheet.docx", build_cheatsheet),
        ("Process_Retraining_Guide.docx", build_process_retraining),
        ("New_Hire_Onboarding_Checklist.docx", build_onboarding),
    ]:
        path = OUTPUT_DIR / name
        doc = builder()
        doc.save(str(path))
        size_kb = path.stat().st_size / 1024
        print(f"  Wrote {name}  ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
