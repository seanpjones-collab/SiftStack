"""Generate Zane's Acquisitions SOP as a polished Word document.

Uses python-docx to produce a printable, scannable .docx with proper
headings, tables, and bullets. Schedule-aware per-day routines based on
Zane's actual hours.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

REPO = Path(__file__).resolve().parent.parent
OUTPUT_PATH = REPO / "output" / "sift_setup" / "Zane_Acquisitions_SOP.docx"

# ── Style helpers ─────────────────────────────────────────────────────


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 0:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_callout(doc: Document, text: str, *, color: str = "warn") -> None:
    """Add a single-cell shaded table acting as a callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Light Grid Accent 1" if color == "info" else "Light Grid Accent 2"
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    # Data rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(10)


def add_pagebreak(doc: Document) -> None:
    doc.add_page_break()


# ── Document content ─────────────────────────────────────────────────


def build_doc() -> Document:
    doc = Document()

    # Page setup — narrower margins for printability
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Cover / Title
    add_heading(doc, "Zane — Acquisitions SOP", level=0)
    add_para(doc, "Lead Manager + Closer combined", italic=True, size=12)
    add_para(doc, "Specialist (Blueprint C) — Cuyahoga / Summit / Stark, Probate + Foreclosure",
             italic=True, size=11)
    add_para(doc, "Last updated 2026-04-28", italic=True, size=10)
    doc.add_paragraph()

    # The 5 non-negotiables
    add_heading(doc, "The 5 Non-Negotiable Rules", level=1)
    add_callout(doc,
                "These are the rules that break the entire system if you skip them. "
                "If something on the list below conflicts with the rest of this SOP, the rule wins.",
                color="warn")
    doc.add_paragraph()
    rules = [
        ("Every lead must have a next step.",
         "No exceptions. If a lead exists in Sift without a scheduled task, it's already dying. "
         "Set the next-step task before you move to the next call."),
        ("Call back within 1 minute on any new inbound inquiry.",
         "Harvard data: 400% close rate uplift vs 5+ minutes. Drop everything for new inbound."),
        ("Update Property Status after every contact.",
         "STABM: Status Accuracy is Foundational. Wrong status fires the wrong sequences. "
         "Status update is part of the call, not a 'later' task."),
        ("Calls between 8 AM and 9 PM only.",
         "TCPA compliance window. Even on Saturday on-call, you can't dial outside this window."),
        ("Empathy first, transactional second.",
         "Probate leads are grieving. Foreclosure leads are being sued. The 'Are you interested "
         "in selling?' opener gets you hung up on. Use the niche-specific scripts in this SOP."),
    ]
    for title, body in rules:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(title + " ")
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run = p.add_run(body)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    # Daily routine - per day
    add_pagebreak(doc)
    add_heading(doc, "Your Daily Routine (Per Day, Based on Your Schedule)", level=1)
    add_para(doc,
             "Each day below is structured around your actual work hours. "
             "Targets and Sift presets are tuned to the hours available. "
             "If you finish a block early, run the next preset in sequence — never sit idle "
             "while there are records in queue.",
             italic=True)

    # Monday
    add_heading(doc, "Monday — 2:00 PM – 8:00 PM (6 hours)", level=2)
    add_para(doc, "Daily target: 8–12 high-quality attempts.")
    monday = [
        ("2:00 – 2:30 PM", "Start of day",
         "Open Sift. Review Tasks tab — anything due today? Run preset "
         "“00 Niche Sequential Marketing → 02. Ready to Call” for fresh FTM. "
         "Run “02 Podio Legacy Pipeline → 00. Hot Active Calls” for hot legacy follow-ups."),
        ("2:30 – 5:00 PM", "Call block 1 (2.5 hrs)",
         "Priority: Hot temp + 1–2-day cadence first → Fresh FTM → Warm Active → Cold Active. "
         "smrtPhone click-to-dial. After every call: status + next-step task."),
        ("5:00 – 5:30 PM", "Admin break",
         "Eat. Reply to email/SMS replies that came in earlier today. Run preset "
         "“03. FTM Follow Up 1” for second-attempt records."),
        ("5:30 – 7:30 PM", "Call block 2 (2 hrs)",
         "Continue priority order. Make offers on qualified leads (passing all 4 Pillars). "
         "Schedule appointments where motivation is high but qualification is incomplete."),
        ("7:30 – 8:00 PM", "End of day",
         "Final pass: every record you touched today has correct status + next-step task. "
         "Send daily KPI to Sean: attempts, RPC count, offers made, contracts."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], monday)

    # Tuesday
    add_heading(doc, "Tuesday — 9:00 AM – 6:00 PM (9 hours)", level=2)
    add_para(doc, "Daily target: 15–20 attempts. Full-day rhythm.")
    tuesday = [
        ("9:00 – 9:30 AM", "Start of day",
         "Open Sift. Review Tasks. Run “02. Ready to Call” + “00. Hot Active Calls (Legacy).” "
         "Note total queue count — flag to Sean if >40."),
        ("9:30 AM – 12:00 PM", "Call block 1 (2.5 hrs)",
         "Hot leads first. Target 8–12 attempts in this block. Update status + next-step task after each call."),
        ("12:00 – 1:00 PM", "Lunch",
         "Eat. Reply to texts/emails that came in mid-morning."),
        ("1:00 – 2:00 PM", "Admin + Follow Ups",
         "Run “03. FTM Follow Up 1.” Work through follow-up tasks due today."),
        ("2:00 – 5:30 PM", "Call block 2 (3.5 hrs)",
         "Continue priority order. Make offers on qualified leads. "
         "Push on appointments and offer follow-ups."),
        ("5:30 – 6:00 PM", "End of day",
         "Final pass on statuses + tasks. Send daily KPI to Sean."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], tuesday)

    # Wednesday
    add_heading(doc, "Wednesday — 9:00 AM – 1:00 PM, then 3:00 PM – 8:00 PM (9 hours, split)", level=2)
    add_para(doc, "Daily target: 12–18 attempts. Two halves with a midday break.")
    wednesday = [
        ("9:00 – 9:30 AM", "Start of day",
         "Open Sift. Review Tasks. Run morning presets."),
        ("9:30 AM – 1:00 PM", "Call block 1 (3.5 hrs)",
         "Hot/fresh priority. Target 8–10 attempts. Status + next step after each call."),
        ("1:00 – 3:00 PM", "BREAK (off the clock)",
         "Off the clock. Don't dial. Don't update Sift. Recharge."),
        ("3:00 – 3:15 PM", "Re-orient",
         "Quick scan: any inbound calls/SMS during the break that need callback within 1 minute? "
         "Otherwise, review afternoon tasks."),
        ("3:15 – 5:00 PM", "Call block 2 (1.75 hrs)",
         "Follow-up and second-attempt focus. Run “03–05. FTM Follow Up 1/2/3” presets."),
        ("5:00 – 5:30 PM", "Admin",
         "Reply to email/SMS. Quick triage."),
        ("5:30 – 7:30 PM", "Call block 3 (2 hrs)",
         "Evening attempts — best hours to catch working sellers at home. "
         "Push offers and appointments."),
        ("7:30 – 8:00 PM", "End of day",
         "Final pass + KPI to Sean."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], wednesday)

    # Thursday
    add_heading(doc, "Thursday — 9:00 AM – 1:00 PM, then 3:00 PM – 8:00 PM (9 hours, split)", level=2)
    add_para(doc, "Same structure as Wednesday. Daily target: 12–18 attempts.")
    add_para(doc, "Follow Wednesday's table above — schedule is identical.", italic=True)

    # Friday
    add_heading(doc, "Friday — 9:00 AM – 5:00 PM (8 hours)", level=2)
    add_para(doc, "Daily target: 12–18 attempts. Wrap up the week clean.")
    friday = [
        ("9:00 – 9:30 AM", "Start of day",
         "Open Sift. Review Tasks. Run morning presets."),
        ("9:30 AM – 12:00 PM", "Call block 1 (2.5 hrs)",
         "Hot/fresh priority. Target 8–10 attempts."),
        ("12:00 – 12:30 PM", "Lunch / SMS triage",
         "Quick lunch. Reply to texts/emails. Note anything that needs a Friday-afternoon callback."),
        ("12:30 – 4:30 PM", "Call block 2 (4 hrs)",
         "Long block — push hard on follow-ups, offers, appointment confirms for next week. "
         "Friday afternoon is when sellers commit or push to next week — get specific."),
        ("4:30 – 5:00 PM", "End of day + week-end reporting",
         "Final status pass. Send Sean: weekly KPIs (total attempts, RPC, offers, "
         "contracts, conversion ratios). Flag anything that needs Sean's attention "
         "over the weekend."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], friday)

    # Saturday
    add_heading(doc, "Saturday — 9:00 AM – 1:00 PM + on-call (4 hours base)", level=2)
    add_para(doc, "Daily target: 5–8 attempts during base hours, plus all inbound.")
    saturday = [
        ("9:00 – 9:30 AM", "Start of day",
         "Saturday queue is mostly: hot leads (maintain 1–2-day cadence), inbound that "
         "came in overnight, and appointment confirms for the coming week."),
        ("9:30 AM – 12:30 PM", "Inbound focus + hot lead callbacks",
         "Saturdays are good for catching W-2 sellers home. Hit your hot list. "
         "Don't blast cold — Saturday is for high-quality outreach only."),
        ("12:30 – 1:00 PM", "End of base shift + on-call handoff",
         "Final status pass. Confirm Sift is up to date for Sean's weekend visibility. "
         "Switch to on-call mode."),
        ("After 1:00 PM", "ON-CALL",
         "Respond to inbound calls + SMS as they come in. New inbound = call back within 1 minute, "
         "even on weekend. Anything that comes in via web form or PPL provider gets the same treatment. "
         "If no inbound, no action needed."),
    ]
    add_table(doc, ["Time", "Block", "What to do"], saturday)

    # Sunday
    add_heading(doc, "Sunday — Off", level=2)
    add_para(doc, "Recharge. Don't dial. Don't update Sift unless something urgent comes through. "
                  "Inbound + emergency only.", italic=True)

    # 4 Pillars
    add_pagebreak(doc)
    add_heading(doc, "The 4 Pillars Qualification", level=1)
    add_para(doc, "Apply on EVERY contact. A lead is qualified when you have answers to all four. "
                  "If 3 of 4 → follow up. If 2 or fewer → cold lead.")
    pillars = [
        ("Reason", "Why are they thinking about it?",
         "Specific event: probate, divorce, job loss, tired landlord, code violation"),
        ("Timeline", "When do they need this resolved?",
         "“ASAP” / “within 30 days” / before auction date / before next mortgage payment"),
        ("Condition", "What's the property like?",
         "Honest about needing work / vacant / inherited as-is / hasn't been maintained"),
        ("Price", "What number are they thinking?",
         "Realistic range OR “open to offers” OR an explicit need (not asking price)"),
    ]
    add_table(doc, ["Pillar", "What to ask", "Hot signal"], pillars)
    doc.add_paragraph()
    add_para(doc, "When all 4 are answered → make the offer that day. Don't slow-play hot leads.",
             bold=True)
    add_para(doc, "If 1–2 are missing → set a next-step task to dig deeper on those specific gaps.")

    # Probate script
    add_pagebreak(doc)
    add_heading(doc, "Probate Script — Opening + Conversation Flow", level=1)
    add_callout(doc,
                "Context: They likely just lost a parent or close family member. They are NOT in "
                "“selling mode” yet. Lead with empathy or you get hung up on.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "The Opener (memorize this)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '"Hi, is this [PR/Executor first name]? My name is Zane with [company]. '
        'First — I noticed you\'re handling [decedent\'s first name]\'s estate, '
        'and I just wanted to say I\'m sorry for your loss. I work with families '
        'navigating inherited properties in Cuyahoga County, and I reach out '
        'because most people aren\'t sure what their options are. Do you have '
        'a couple of minutes? I won\'t take long."'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()
    add_para(doc, "Why this works:", bold=True)
    add_bullet(doc, "Acknowledges the loss before any pitch")
    add_bullet(doc, "Frames you as a resource, not a buyer")
    add_bullet(doc, "Gives them an explicit “couple of minutes” exit they can refuse easily")

    add_heading(doc, "If they engage", level=2)
    add_para(doc, "Don't ask “are you selling.” Ask what they're dealing with:")
    p = doc.add_paragraph()
    run = p.add_run(
        '"What\'s the situation with the property — are you keeping it, renting '
        'it out, or thinking about selling once probate clears?"'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()
    add_para(doc, "Listen. Most will fall into one of three buckets:")
    add_bullet(doc, '"We don\'t know yet" → Education mode. Explain probate timeline + options. No pressure. Set 30-day follow-up.')
    add_bullet(doc, '"We want to keep it" → "Got it. If anything changes, mind if I reach out in a few months?" Tag, schedule 90-day follow-up.')
    add_bullet(doc, '"We\'re going to sell" → 4 Pillars on the spot. These are most likely candidates for an offer.')

    add_heading(doc, "Discovery questions if they're going to sell (in this order)", level=2)
    add_numbered(doc, '"Has the probate court appointed you as the executor / personal rep yet?" — determines authority to sell')
    add_numbered(doc, '"Has anyone else looked at the property — agents, family, other investors?" — competition check')
    add_numbered(doc, '"What kind of shape is it in? Honestly — anything you\'ve been putting off?" — condition check')
    add_numbered(doc, '"If everything was easy and you got a fair price, what number would feel right to you?" — price expectation, NOT "asking price"')

    add_heading(doc, "Make the offer (when 4 Pillars confirmed)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '"Based on what you told me — [property condition + their situation], I think '
        'I can help. I can make you a cash offer with no agent fees, no repairs, and '
        'we can close as fast as the court will let us. My number would be in the '
        '[$X to $Y] range. If that\'s in the ballpark, we can talk specifics. If '
        'it\'s too low, no hard feelings — I just want to be straight with you."'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, "NEVER say to a probate lead", level=2)
    add_bullet(doc, '"Are you interested in selling?" — transactional in a grief moment, hang-up trigger')
    add_bullet(doc, '"Sorry to bother you at a difficult time" — you\'re not bothering them; don\'t apologize for existing')
    add_bullet(doc, '"I see your dad just passed" — using "dad" feels invasive. Use the decedent\'s first name.')
    add_bullet(doc, '"I can give you a great deal" — skeptical, salesy')

    # Foreclosure script
    add_pagebreak(doc)
    add_heading(doc, "Foreclosure Script — Opening + Conversation Flow", level=1)
    add_callout(doc,
                "Context: They're being sued by their lender. Mailbox stuffed with notices, "
                "calls from creditors, possibly facing losing their home. They've heard 100 "
                "“we buy houses” pitches. Asking “are you selling” gets you hung up on or "
                "sworn at. Lead with help, not pressure.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "The Opener", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '"Hi, is this [first name]? Zane here. I\'m reaching out because I work '
        'with homeowners in [Cuyahoga / Summit / Stark] who are dealing with the '
        'foreclosure process, and I noticed your situation. Look — I\'m not '
        'calling to pitch you anything right now. I just want to make sure you '
        'know what your options are, because most people don\'t get told. Got '
        'a couple minutes?"'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()
    add_para(doc, "Why this works:", bold=True)
    add_bullet(doc, '"I\'m not calling to pitch you" disarms the "this is a sales call" reflex')
    add_bullet(doc, '"Make sure you know your options" frames you as a resource')
    add_bullet(doc, '"Most people don\'t get told" creates a small information gap they want to close')

    add_heading(doc, "If they engage", level=2)
    add_para(doc, "Most will be defensive at first. Don't push selling. Ask what they want:")
    p = doc.add_paragraph()
    run = p.add_run(
        '"First question — are you trying to keep the house and get caught up, '
        'or are you ready to be done with it?"'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()
    add_para(doc, "Listen for the answer:")
    add_bullet(doc, '"Trying to keep it" → "Have you talked to a HUD-approved housing counselor? They\'re free and can help you negotiate with the bank. I can send you the link." Send it. Tag with `keeping-house`. Follow up in 30 days. Some who try to keep it can\'t — they remember who was honest with them.')
    add_bullet(doc, '"Ready to be done" → 4 Pillars. Likely qualified.')
    add_bullet(doc, '"I don\'t know" → Discovery mode. Don\'t push.')

    add_heading(doc, "Discovery questions for “ready to be done”", level=2)
    add_numbered(doc, '"When\'s the auction date — do you have one set?" — timing pressure')
    add_numbered(doc, '"How much do you owe on it, roughly?" — determines if there\'s equity')
    add_numbered(doc, '"What\'s the place look like inside — any work it needs?" — condition')
    add_numbered(doc, '"What would you NEED to walk away whole and start fresh?" — price framed as their NEEDS, NOT "asking price"')

    add_heading(doc, "Make the offer (when qualified)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '"Based on what you told me, here\'s what I can do — I can pay off your '
        'loan and put [$X] in your pocket. No agent fees, no repairs, you take '
        'what you want and leave the rest. I can close before the auction date. '
        'Does that get you what you need?"'
    )
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, "NEVER say to a foreclosure lead", level=2)
    add_bullet(doc, '"I see you\'re in foreclosure" — don\'t lead with their shame')
    add_bullet(doc, '"Don\'t worry, I can save you" — savior complex, patronizing')
    add_bullet(doc, '"I can give you a great deal" — you\'re getting their house cheap; don\'t gloat')
    add_bullet(doc, '"How much are you asking?" — they have no asking price; ask what they NEED')

    # Status update playbook
    add_pagebreak(doc)
    add_heading(doc, "Status Update Playbook (after every call)", level=1)
    add_para(doc, "Use this exact mapping when you set Property Status in Sift. No improvising.")

    statuses = [
        ("Right party, qualified, will sell", "Hot Lead", "Offer call in 1–2 days"),
        ("Right party, interested but not ready", "Warm Lead", "Follow-up in 15 days"),
        ("Right party, not motivated yet (probate)", "Cold Lead", "Re-engage in 45 days"),
        ("Right party, not motivated yet (foreclosure / auction-driven)", "Cold Lead", "Re-engage in 30 days"),
        ("Right party, not motivated yet (general)", "Cold Lead", "Re-engage in 90 days"),
        ("Right party, hard no", "Not Interested", "Quarterly re-engage in 90 days"),
        ("Wrong number / dead number", "(don't change status — Sift's phone status auto-updates)", "Move to next number"),
        ("Right party, asked to be left alone", "Dead Lead + DNC tag", "Don't call again"),
        ("Already sold to someone else", "Sold", "Done"),
        ("Hostile / threatening", "Dead Lead + note in message board", "Done"),
    ]
    add_table(doc, ["Outcome", "Property Status", "Next-step task"], statuses)
    doc.add_paragraph()
    add_callout(doc,
                "Golden rule: every lead leaves the call with a scheduled task. "
                "Even a Dead Lead gets a “verify in 6 months” if there's any chance "
                "circumstances change.",
                color="warn")

    # KPIs
    add_pagebreak(doc)
    add_heading(doc, "Your KPIs (what Sean watches weekly)", level=1)
    kpis = [
        ("Daily attempts", "Per schedule above (8–20 depending on day length)"),
        ("Right Party Contacts (RPC) per attempt", "1 in 32 baseline; with Trestle scoring, target 1 in 10"),
        ("Conversations to appointment", "5:1"),
        ("Appointments to offer", "1.5:1"),
        ("Offers to contract", "4:1"),
        ("Task completion rate", "95% or higher"),
        ("Tasks overdue", "Less than 5%"),
    ]
    add_table(doc, ["Metric", "Target"], kpis)
    doc.add_paragraph()
    add_para(doc,
             "Trestle phone scoring (the 5-tier system) drives a 4.75× connect-rate uplift "
             "from baseline. Always dial 81–100 score numbers first, then 61–80. "
             "Skip anything 40 or below unless every other lead is exhausted.",
             italic=True)

    # When to escalate
    add_heading(doc, "When to escalate to Sean", level=1)
    add_bullet(doc, "Property under contract — handoff to Sean for transaction coordination")
    add_bullet(doc, "Seller wants to negotiate beyond your authority — pull Sean into the call")
    add_bullet(doc, "Lead has unusual situation (out-of-state heirs, divorce mid-sale, title issues) — Sean handles or routes to data manager for Deep Prospecting")
    add_bullet(doc, "Anything legally weird (active eviction, code violations, tenant disputes) — escalate")
    add_bullet(doc, "Inbound that looks like it might be press/competitor/scammer — flag and ask")

    # What works / doesn't
    add_pagebreak(doc)
    add_heading(doc, "Things that work (per Ty's curriculum)", level=1)
    add_bullet(doc, "Empathy openers in distress niches — probate/foreclosure leads convert 2–3× better with empathy-first")
    add_bullet(doc, '"I\'m not pitching you" disarm — flips a defensive caller into curious')
    add_bullet(doc, "Asking for THEIR number, not the asking price — “what would you need to walk away whole” is psychologically different from “what are you asking”")
    add_bullet(doc, "The 1-minute callback rule — single highest-ROI behavior change")
    add_bullet(doc, "Updating status immediately after each call — broken statuses break the entire system")

    add_heading(doc, "Things that don't (per Ty)", level=1)
    add_bullet(doc, 'Generic "are you interested in selling" openers on distressed sellers')
    add_bullet(doc, "Skipping the empathy step to save time")
    add_bullet(doc, 'Marking leads as "Hot" without all 4 Pillars confirmed')
    add_bullet(doc, "Leaving any record without a next-step task")
    add_bullet(doc, "Power-dialing niche FTM records (these are high-value first-to-market — they deserve personal touch)")

    # Footer
    doc.add_paragraph()
    add_para(doc,
             "Questions, situations not covered here, or anything that feels off — "
             "ask Sean. This SOP is a starting point, not a cage.",
             italic=True, size=10)

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(str(OUTPUT_PATH))
    size_kb = OUTPUT_PATH.stat().st_size / 1024
    print(f"Wrote {OUTPUT_PATH}")
    print(f"Size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
