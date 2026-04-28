"""Generate Zane's companion training docs as a polished Word document.

Includes:
  - Objection handling (probate + foreclosure)
  - Voicemail scripts (Day 1 / Day 2 / Day 3 per niche)
  - SMS follow-up templates (Day 1 / Day 2 / Day 3 per niche)
  - Email templates (Day 1 + follow-up per niche)
  - Tone + style rules
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches, RGBColor

REPO = Path(__file__).resolve().parent.parent
OUTPUT_PATH = REPO / "output" / "sift_setup" / "Zane_Companion_Training.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 0:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def add_para(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_callout(doc, text, color="warn"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Light Grid Accent 1" if color == "info" else "Light Grid Accent 2"
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True


def add_quote(doc, text, *, label=None):
    """Add a quoted/scripted line in italic with optional bold label above."""
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_objection_pair(doc, objection, response):
    """Two-column table-like layout: objection on left, scripted response on right."""
    p = doc.add_paragraph()
    run = p.add_run("They say: ")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run = p.add_run(f'"{objection}"')
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("You say: ")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run = p.add_run(f'"{response}"')
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()


def add_pagebreak(doc):
    doc.add_page_break()


def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Cover
    add_heading(doc, "Zane — Companion Training", level=0)
    add_para(doc, "Objection Handling, Voicemails, SMS, Email Templates", italic=True, size=12)
    add_para(doc, "Companion to: Zane Acquisitions SOP", italic=True, size=11)
    add_para(doc, "Last updated 2026-04-28", italic=True, size=10)
    doc.add_paragraph()

    add_callout(doc,
                "Use this with the SOP, not instead of it. The SOP tells you WHEN to call, "
                "WHAT preset, WHICH status. This document tells you HOW to handle the call once "
                "the phone rings.",
                color="warn")
    doc.add_paragraph()

    # Tone rules
    add_heading(doc, "Tone Rules (read once, internalize)", level=1)
    add_bullet(doc, "Calm, low-pressure. Distressed sellers are sensitive to anyone who sounds urgent or hungry.")
    add_bullet(doc, "Speak slower than you naturally would. Drops anxiety on their end.")
    add_bullet(doc, "Pause after they finish a sentence. Most people fill silence with more information — that's where the qualifying details come from.")
    add_bullet(doc, "Mirror their language. If they say “the house,” don't say “the property.” If they say “my mom,” don't say “the deceased.”")
    add_bullet(doc, "Never argue. If they push back, agree, ask one more question, then back off.")
    add_bullet(doc, "Smile when you talk. They can hear it. Even on a script you've read 100 times.")

    # ─────────────────────────────────────────────────────────────────
    # OBJECTIONS — PROBATE
    # ─────────────────────────────────────────────────────────────────
    add_pagebreak(doc)
    add_heading(doc, "Objection Handling — Probate Leads", level=1)
    add_para(doc,
             "These are the most common objections you'll hear from probate leads, and the "
             "scripted response that keeps the conversation alive without pushing.",
             italic=True)

    add_objection_pair(
        doc,
        "We're not selling.",
        "Totally understand — most families don't think about it right after. Mind if I "
        "send you a free probate guide so when you're ready to think about it, you have "
        "the info? No follow-up unless you ask."
    )
    add_objection_pair(
        doc,
        "We're going to keep it in the family.",
        "That's great. Most people end up keeping it. If anything changes — kids don't "
        "want it, taxes get heavy, repairs pile up — would it be okay if I checked in "
        "once a year? Just a quick text."
    )
    add_objection_pair(
        doc,
        "We have an agent already.",
        "Oh good. Just so you know your options — agents typically need 30 to 90 days "
        "and you'd need to deal with showings, repairs, and inspections. We can usually "
        "close in 2 to 3 weeks cash if that ever sounds appealing. Otherwise, I'm out "
        "of your hair."
    )
    add_objection_pair(
        doc,
        "I'm in the middle of probate. I can't sell yet.",
        "Right, the court has to appoint you first. Has that happened? Usually takes "
        "30 to 60 days. If you want, I can put your file aside and check back in a month — "
        "no obligation, just so we're ready when you are."
    )
    add_objection_pair(
        doc,
        "What's it worth?",
        "Honest answer — depends on condition. If you can spare 5 minutes, I can give "
        "you a ballpark over the phone, or I can swing by and walk through it. "
        "Whichever you prefer."
    )
    add_objection_pair(
        doc,
        "How did you find out about this?",
        "Public record — the probate filing in [county] court. I track those because "
        "most families don't realize how much help is available."
    )
    add_objection_pair(
        doc,
        "Stop calling me.",
        "Got it, I'll take you off the list right now. If anything changes, my number "
        "is [number]. Take care."
    )
    add_objection_pair(
        doc,
        "Are you a scammer?",
        "Fair question. I'm a local investor in [city] — you can verify the company at "
        "[website]. Take your time. I'd rather you check first than feel pressured."
    )

    # ─────────────────────────────────────────────────────────────────
    # OBJECTIONS — FORECLOSURE
    # ─────────────────────────────────────────────────────────────────
    add_pagebreak(doc)
    add_heading(doc, "Objection Handling — Foreclosure Leads", level=1)
    add_para(doc,
             "Foreclosure leads are defensive by default. Every response below is "
             "designed to lower the temperature and offer real options before any sales pitch.",
             italic=True)

    add_objection_pair(
        doc,
        "I'm not selling.",
        "Got it. Just so you know — there are a few options most folks don't realize: "
        "HUD counseling (free), short sale, deed in lieu, and yes, a cash sale. Mind if "
        "I send a one-page summary so you have it for reference? No pressure."
    )
    add_objection_pair(
        doc,
        "I'm working with my bank.",
        "Good. Are they offering you a modification or asking you to sell? I ask because "
        "I see how lenders move every day, and I can usually tell if they're being "
        "straight with you. Happy to give you my honest take if it helps."
    )
    add_objection_pair(
        doc,
        "I have an attorney.",
        "Smart move. While they're working that, do you have a backup plan in case it "
        "doesn't work out? Sometimes the legal side takes longer than the auction date "
        "and people get caught. I'm not saying that's you — just want to make sure "
        "you've thought it through."
    )
    add_objection_pair(
        doc,
        "I'll figure it out myself.",
        "Totally fair. If I can ask one thing — do you know your auction date? That's "
        "the clock that matters. Once I know that, I can tell you whether you've got "
        "time to figure it out, or whether you're cutting it close."
    )
    add_objection_pair(
        doc,
        "How did you get my information?",
        "Public record — the foreclosure filing in [county] court. I track those "
        "because most homeowners in that situation don't get told their options."
    )
    add_objection_pair(
        doc,
        "Stop calling me.",
        "Done. I'll take you off the list. If your situation changes, my number is "
        "[number]. Hope it works out for you."
    )
    add_objection_pair(
        doc,
        "Are you one of those investors trying to lowball me?",
        "Fair to ask. I'm not going to lowball you — what I can pay depends on what "
        "you owe and what condition the place is in. If my number doesn't work for "
        "you, no hard feelings, you walk away. I just want to be a real option, not a "
        "vulture."
    )
    add_objection_pair(
        doc,
        "I just need a few more weeks.",
        "Okay. Do you have a specific plan in those few weeks — bank work-out, family "
        "loan, refinance? I ask because if it doesn't come through, the auction date "
        "doesn't move. If you want to put a backup in place just in case, I can have "
        "an offer ready to go within 24 hours."
    )

    # ─────────────────────────────────────────────────────────────────
    # VOICEMAILS
    # ─────────────────────────────────────────────────────────────────
    add_pagebreak(doc)
    add_heading(doc, "Voicemail Scripts", level=1)
    add_callout(doc,
                "Voicemails follow the 3-day cadence. Each VM is shorter than the last — "
                "by Day 3 you're respecting their silence. Don't keep escalating after Day 3 "
                "unless you have a new reason to call (auction date, court date, etc.).",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "Probate Voicemails", level=2)
    add_quote(doc,
              "Hi [first name], this is Zane with [company]. I'm reaching out about "
              "[decedent first name]'s estate — sorry for your loss. I help families "
              "in [county] navigate inherited properties. No pressure on this call — I "
              "just want to make sure you know your options when you're ready. My "
              "number is [number]. Take care.",
              label="VM 1 — Day 1 (~25 seconds)")
    add_quote(doc,
              "Hi [first name], Zane again. Just wanted to make myself available. If "
              "you're sorting through what to do with the property and want a second "
              "opinion or just info on the process, I'm happy to help — no pressure, "
              "no pitch. [number]. Talk soon.",
              label="VM 2 — Day 2 (~20 seconds)")
    add_quote(doc,
              "Hi [first name], last try from me. If now's not the right time, totally "
              "understand. If you ever want to talk about [property address] or "
              "anything related to [decedent first name]'s estate, I'm at [number]. "
              "Take care.",
              label="VM 3 — Day 3 (~15 seconds)")

    add_heading(doc, "Foreclosure Voicemails", level=2)
    add_quote(doc,
              "Hi [first name], Zane here. I work with homeowners going through the "
              "foreclosure process in [county]. I'm not calling to sell you anything "
              "— I just want to make sure you know your options. There's more than "
              "people realize, and most of them are free. My number is [number]. "
              "Take care.",
              label="VM 1 — Day 1 (~25 seconds)")
    add_quote(doc,
              "Hi [first name], Zane again. Quick one — if you're trying to keep the "
              "house, there's free HUD help. If you're ready to walk away, there are "
              "options too. Either way, I'm happy to walk through it with you. "
              "[number].",
              label="VM 2 — Day 2 (~20 seconds)")
    add_quote(doc,
              "Hi [first name], one more try. If you've got an auction date coming "
              "up, the clock matters. If you want help understanding your options "
              "before then, I'm at [number]. No pressure either way.",
              label="VM 3 — Day 3 (~15 seconds)")

    # ─────────────────────────────────────────────────────────────────
    # SMS FOLLOW-UPS
    # ─────────────────────────────────────────────────────────────────
    add_pagebreak(doc)
    add_heading(doc, "SMS Follow-Up Templates", level=1)
    add_callout(doc,
                "SMS goes after the call attempt + voicemail on the same day. Short, "
                "specific, no marketing language. Always sign off with your name.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "Probate SMS Sequence", level=2)
    add_quote(doc,
              "Hi [first name], this is Zane with [company]. I left you a voicemail "
              "about [decedent first name]'s estate. I help families with inherited "
              "properties — no pressure, just info if you want it. Reply anytime.",
              label="SMS 1 — Day 1, after VM")
    add_quote(doc,
              "Hi [first name], following up. If you're sorting through what to do "
              "with [property address], happy to walk through options no charge. "
              "-Zane",
              label="SMS 2 — Day 2")
    add_quote(doc,
              "Hi [first name], last text from me unless you reach out. My number "
              "is [number] if you ever want to talk about [property address]. "
              "Take care. -Zane",
              label="SMS 3 — Day 3")

    add_heading(doc, "Foreclosure SMS Sequence", level=2)
    add_quote(doc,
              "Hi [first name], Zane here. I help homeowners in [county] dealing "
              "with foreclosure. Not pitching anything — just options. Reply if "
              "you want to know what they are.",
              label="SMS 1 — Day 1, after VM")
    add_quote(doc,
              "Hi [first name], if you're trying to keep the house, free HUD help "
              "exists. If you're ready to walk away, I can usually close before any "
              "auction date. Either way I'm here. -Zane",
              label="SMS 2 — Day 2")
    add_quote(doc,
              "Hi [first name], last text. Number's [number] if you change your "
              "mind. Wishing you the best. -Zane",
              label="SMS 3 — Day 3")

    add_heading(doc, "Reply handling", level=2)
    add_bullet(doc, '"Stop" / "Don\'t text me" / "F off" → Reply with "Got it, taking you off. Take care." Then mark Property Status = Dead Lead, add DNC tag.')
    add_bullet(doc, '"Who is this?" / "How did you get my number?" → "Public record — the [foreclosure / probate] filing in [county] court. Happy to send my contact info if you want to verify."')
    add_bullet(doc, '"Maybe / I\'m thinking about it" → "Happy to chat whenever. When\'s a good time to call?" Schedule the callback as a Sift task.')
    add_bullet(doc, '"How much would you pay?" → "Depends on condition + what you owe. Easiest is a quick 5-min call — got 5 min today?" Schedule the call.')

    # ─────────────────────────────────────────────────────────────────
    # EMAIL TEMPLATES
    # ─────────────────────────────────────────────────────────────────
    add_pagebreak(doc)
    add_heading(doc, "Email Templates", level=1)
    add_callout(doc,
                "Email is the FIRST channel in the Pendulum (Email → SMS → Call → Mail → DP). "
                "It's the cheapest touch and creates a paper trail. Use it as the lead-in to "
                "the call sequence, not as a substitute for it.",
                color="warn")
    doc.add_paragraph()

    add_heading(doc, "Probate Email — Day 1 (initial outreach)", level=2)
    add_para(doc, "Subject: A quick note about [decedent first name]'s estate", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              "Hi [first name],\n\n"
              "I'm Zane with [company] in [city]. First — I'm sorry for your loss. "
              "I came across [decedent first name]'s estate in the probate court "
              "records and wanted to reach out.\n\n"
              "I work with families in [county] navigating inherited properties. "
              "Most of the time, families aren't sure what their options are: keep "
              "and rent, sell traditionally, or sell as-is for cash. There's no "
              "right answer — it depends on what you and your family want.\n\n"
              "If you'd like to talk through it (no pitch, just information), my "
              "direct line is [number]. I'm also happy to send over a one-page "
              "summary of the options if that's easier.\n\n"
              "Take care,\nZane")

    add_heading(doc, "Probate Email — Day 7 follow-up", level=2)
    add_para(doc, "Subject: RE: A quick note about [decedent first name]'s estate", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              "Hi [first name],\n\n"
              "Following up on my earlier note. I know there's a lot on your plate "
              "right now, so no pressure.\n\n"
              "If [property address] is something you're starting to think about, "
              "I can give you a no-obligation cash offer in 48 hours. If you'd "
              "rather list it traditionally or hold onto it, I can point you to "
              "people who can help with that too.\n\n"
              "Either way, my direct line is [number].\n\n"
              "Take care,\nZane")

    add_heading(doc, "Foreclosure Email — Day 1 (initial outreach)", level=2)
    add_para(doc, "Subject: Quick note about [property address]", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              "Hi [first name],\n\n"
              "I'm Zane with [company] in [city]. I came across the foreclosure "
              "filing for [property address] in [county] court records and wanted "
              "to reach out — not to pitch you anything, but to make sure you "
              "know what your options are.\n\n"
              "Most homeowners in this situation don't get told that there are "
              "several paths forward:\n\n"
              "1. HUD-approved housing counselors (free) help you negotiate with "
              "the bank for a loan modification\n"
              "2. Short sale — your bank approves a sale for less than what you "
              "owe\n"
              "3. Deed in lieu — give the property back to the bank without going "
              "through auction\n"
              "4. Cash sale — sell to an investor (like me) before the auction "
              "and walk away with cash in your pocket\n\n"
              "Each one has trade-offs. Happy to walk through them with you if "
              "you'd like — no pitch, just information. My direct line is "
              "[number].\n\n"
              "Either way, please don't let the auction date sneak up on you. "
              "Take care,\nZane")

    add_heading(doc, "Foreclosure Email — Day 7 follow-up", level=2)
    add_para(doc, "Subject: RE: Quick note about [property address]", bold=True)
    doc.add_paragraph()
    add_para(doc, "Body:", bold=True)
    add_quote(doc,
              "Hi [first name],\n\n"
              "Following up. If you've already figured out a path, that's great — "
              "I hope it works out.\n\n"
              "If you haven't, the clock is the thing that matters most here. "
              "Auction dates don't move because you're working on it.\n\n"
              "If you want to put a backup option in place — even just to know "
              "the number — I can have a written cash offer ready in 24 hours. "
              "No obligation. You compare it to whatever else you're working on.\n\n"
              "Direct line: [number].\n\n"
              "Take care,\nZane")

    # ─────────────────────────────────────────────────────────────────
    # WHAT TO DO AT EACH OUTCOME
    # ─────────────────────────────────────────────────────────────────
    add_pagebreak(doc)
    add_heading(doc, "Quick Reference — What to Do at Each Call Outcome", level=1)

    add_heading(doc, "They picked up + had a real conversation", level=2)
    add_bullet(doc, "Apply 4 Pillars (Reason / Timeline / Condition / Price) — see SOP")
    add_bullet(doc, "Update Property Status to match outcome (Hot / Warm / Cold / Not Interested)")
    add_bullet(doc, "Set the next-step task at the cadence: Hot 1-2 days, Warm 15 days, Cold 30/45/90 days by niche")
    add_bullet(doc, "Send same-day SMS confirming what was discussed")

    add_heading(doc, "They picked up but pushed back", level=2)
    add_bullet(doc, "Use the objection-handling response from this doc")
    add_bullet(doc, "If they end the call without engaging — Property Status = Cold Lead, follow up on niche cadence")
    add_bullet(doc, "If they explicitly said “stop calling” — Dead Lead + DNC tag, no further outreach")

    add_heading(doc, "Voicemail (no pickup)", level=2)
    add_bullet(doc, "Leave the niche-appropriate VM (Day 1 / 2 / 3 script)")
    add_bullet(doc, "Send the matching SMS within 5 minutes of the VM (one-two punch)")
    add_bullet(doc, "Sift call_attempts auto-increments — preset will surface them on the next day's queue")

    add_heading(doc, "No answer, no voicemail (call dropped)", level=2)
    add_bullet(doc, "Send the SMS only — no VM = no message to leave")
    add_bullet(doc, "Don't mark the phone as wrong unless you've tried 2-3 times across days")

    add_heading(doc, "Wrong number / wrong person", level=2)
    add_bullet(doc, "Apologize politely, hang up — don't try to convert wrong-person calls")
    add_bullet(doc, "Mark that specific phone status as Wrong (per-phone field, not Property Status)")
    add_bullet(doc, "If all phones on the record become Wrong, the Vacant Mailing → DP and No Response DM → DP presets will catch it for the data manager's deep prospecting queue")

    # Footer
    doc.add_paragraph()
    add_para(doc,
             "If you hit a situation not covered here — pause the call if you can, ask Sean, "
             "or schedule a callback. Don't wing it on something legally weird "
             "(active eviction, code violations, tenant disputes, divorces mid-sale).",
             italic=True, size=10)

    return doc


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(str(OUTPUT_PATH))
    size_kb = OUTPUT_PATH.stat().st_size / 1024
    print(f"Wrote {OUTPUT_PATH}")
    print(f"Size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
